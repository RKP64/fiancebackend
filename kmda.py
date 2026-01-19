
import uvicorn
from google import genai
from fastapi import (
    FastAPI,
    UploadFile,
    File,
    HTTPException,
    Depends,
    Form,
    status,
    Request,
    BackgroundTasks,
)
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
import openai
from openai import AzureOpenAI
import json
import os
import tempfile
import traceback
import re
import docx
import html
from docx.enum.text import WD_ALIGN_PARAGRAPH
import httpx  # Using httpx for async requests
from docx import Document
import io
from bs4 import BeautifulSoup, NavigableString
from datetime import datetime, timezone, timedelta
from pymongo import MongoClient
from azure.core.credentials import AzureKeyCredential
from azure.ai.contentsafety import ContentSafetyClient
from azure.ai.contentsafety.models import AnalyzeTextOptions, TextCategory
from azure.search.documents.aio import (
    SearchClient as AsyncSearchClient,
)  # Async search client
from azure.search.documents import SearchClient as SyncSearchClient
from azure.core.credentials import AzureKeyCredential
from azure.search.documents.models import VectorizedQuery
from langchain_openai import AzureOpenAIEmbeddings
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions
from azure.core.exceptions import AzureError
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from dotenv import load_dotenv
import time
import logging
import asyncio
import pandas as pd
import base64
import uuid
import tiktoken  # Added for token counting
import functools  # Added for decorator
import urllib.parse
# ==============================================================================
#  FINANCIAL ENGINE HELPERS (Paste this ABOVE your handle_chat_request function)
# ==============================================================================
import yfinance as yf
import matplotlib
matplotlib.use('Agg') # Essential for backend server
import matplotlib.pyplot as plt
import io
import base64
import time
import re

# --- 1. CACHE ---
_stock_cache = {}
CACHE_DURATION_SECONDS = 300 

def get_cached_stock_info(ticker: str):
    """Fetches stock info with caching."""
    now = time.time()
    if ticker in _stock_cache:
        cached_item = _stock_cache[ticker]
        if now - cached_item["timestamp"] < CACHE_DURATION_SECONDS:
            return cached_item["data"]
    try:
        stock = yf.Ticker(ticker)
        info = stock.info 
        _stock_cache[ticker] = { "data": info, "timestamp": now }
        return info
    except Exception as e:
        print(f"Error fetching {ticker}: {e}")
        return None

# --- 2. EXTRACT TICKERS (The missing function!) ---
async def extract_tickers(query: str) -> list:
    """Maps common company names to NSE tickers."""
    q = query.lower()
    tickers = []
    
    # Add more companies here as needed
    mapping = {
        "tata motors": "TATAMOTORS.NS",
        "maruti": "MARUTI.NS",
        "ola": "OLAELEC.NS", 
        "zomato": "ZOMATO.NS",
        "paytm": "PAYTM.NS",
        "hdfc": "HDFCBANK.NS",
        "reliance": "RELIANCE.NS",
        "ather": "ATHER.NS",
        "tesla": "TSLA",
        "apple": "AAPL"
    }
    
    for name, ticker in mapping.items():
        if name in q:
            tickers.append(ticker)
            
    return list(set(tickers))

# --- 3. GENERATE TABLE ---
def get_financial_table_markdown(tickers: list) -> str:
    if not tickers: return ""
    rows = []
    for t in tickers:
        info = get_cached_stock_info(t)
        if not info: continue
        
        name = info.get('shortName', t)
        price = info.get('currentPrice', 'N/A')
        currency = info.get('currency', 'INR')
        pe = info.get('trailingPE', 'N/A')
        
        mc = info.get('marketCap', 0)
        if mc and isinstance(mc, (int, float)):
            if mc > 1e12: mc_str = f"{mc/1e12:.2f}T"
            elif mc > 1e9: mc_str = f"{mc/1e9:.2f}B"
            else: mc_str = f"{mc/1e6:.2f}M"
        else:
            mc_str = "N/A"

        rows.append(f"| {name} | {currency} {price} | {mc_str} | {pe} |")

    if not rows: return ""
    table = "\n\n### 📈 Market Snapshot\n"
    table += "| Company | Price | Market Cap | P/E Ratio |\n"
    table += "| :--- | :--- | :--- | :--- |\n"
    table += "\n".join(rows)
    table += "\n\n"
    return table

# --- 4. GENERATE VISUALS (TABLE ONLY - NO PNG) ---
def get_stock_chart_markdown(tickers: list) -> str:
    """Generates only the data table, no PNG image."""
    return "" # We already have the table from 'get_financial_table_markdown'

# --- 5. GET CONTEXT FOR LLM (PRO VERSION) ---
def get_llm_financial_context(tickers: list) -> str:
    context = "### CRITICAL INSTRUCTION: LIVE FINANCIAL DATA PROVIDED\n"
    context += "Use this real-time data as your PRIMARY source. Do not apologize if internal docs are missing.\n\n"
    
    has_data = False
    for t in tickers:
        info = get_cached_stock_info(t)
        if info:
            has_data = True
            # -- Extract Core Data --
            currency = info.get('currency', 'INR')
            price = info.get('currentPrice', info.get('regularMarketPrice', 'N/A'))
            
            # Valuation
            pe = info.get('trailingPE', 'N/A')
            f_pe = info.get('forwardPE', 'N/A')
            peg = info.get('pegRatio', 'N/A')
            pb = info.get('priceToBook', 'N/A')
            
            # Fundamentals
            roe = info.get('returnOnEquity', 0) * 100 if info.get('returnOnEquity') else 'N/A'
            rev_growth = info.get('revenueGrowth', 0) * 100 if info.get('revenueGrowth') else 'N/A'
            debt_equity = info.get('debtToEquity', 'N/A')
            margins = info.get('profitMargins', 0) * 100 if info.get('profitMargins') else 'N/A'
            
            # Analyst View
            target_price = info.get('targetMeanPrice', 'N/A')
            recommendation = info.get('recommendationKey', 'N/A').upper()

            # -- Build the Context String --
            context += f"## 📊 Financial Report for {t}\n"
            context += f"* **Price:** {currency} {price}\n"
            context += f"* **Valuation:** P/E: {pe} | PEG: {peg} | P/B: {pb}\n"
            context += f"* **Growth:** Revenue Growth (YoY): {rev_growth}%\n"
            context += f"* **Health:** ROE: {roe}% | Net Margin: {margins}% | Debt/Eq: {debt_equity}\n"
            context += f"* **Analyst Consensus:** {recommendation} (Target: {target_price})\n\n"
            
    return context if has_data else ""

# --- AZURE MONITOR OPENTELEMETRY IMPORTS ---
from azure.monitor.opentelemetry import configure_azure_monitor
from opentelemetry import trace
from opentelemetry.trace import Status, StatusCode
from opentelemetry.context import Context, set_value, get_value, attach, detach
from opentelemetry.instrumentation.fastapi import FastAPIInstrumentor
from opentelemetry.instrumentation.requests import (
    RequestsInstrumentor,
)  # Corrected import if needed
from opentelemetry.instrumentation.logging import LoggingInstrumentor
from opentelemetry.instrumentation.pymongo import PymongoInstrumentor
from opentelemetry.instrumentation.openai import (
    OpenAIInstrumentor as OpenTelemetryOpenAIInstrumentor,
)
from opentelemetry.exporter.otlp.proto.http.trace_exporter import OTLPSpanExporter
from opentelemetry.sdk.resources import Resource
from opentelemetry.sdk.trace import TracerProvider
from opentelemetry.sdk.trace.export import BatchSpanProcessor
from openinference.instrumentation.openai import OpenAIInstrumentor

import functools
from typing import Dict, Any, Optional

# --- SECURITY IMPORTS ---
from jose import JWTError, jwt
from passlib.context import CryptContext
import jwt as pyjwt

# --- Load Environment Variables ---
load_dotenv()

# --- OpenTelemetry Context Variable Keys ---
USER_ID_KEY = "user_id"
USER_ROLE_KEY = "user_role"
USER_QUERY_KEY = "user_query"

# --- Enhanced Logging Setup ---
DEV_MODE = os.environ.get("DEV_MODE", "false").lower() == "true"
log_level = logging.DEBUG if DEV_MODE else logging.INFO

# Create a more detailed formatter
formatter = logging.Formatter(
    "%(asctime)s - %(name)s - %(levelname)s - %(funcName)s:%(lineno)d - %(message)s"
)

# Configure root logger
root_logger = logging.getLogger()
root_logger.setLevel(log_level)

# Remove existing handlers to avoid duplicates
for handler in root_logger.handlers[:]:
    root_logger.removeHandler(handler)

# Create console handler for terminal output
console_handler = logging.StreamHandler()
console_handler.setLevel(log_level)
console_handler.setFormatter(formatter)
root_logger.addHandler(console_handler)

# --- MODIFICATION: Silence noisy SDK loggers in production ---
logging.getLogger("azure.core.pipeline.policies").setLevel(logging.WARNING)
logging.getLogger("openai").setLevel(logging.WARNING)
logging.getLogger("matplotlib").setLevel(logging.WARNING)
logging.getLogger("azure.monitor.opentelemetry.exporter.export._base").setLevel(
    logging.WARNING
)  # Silence transmission logs
logging.getLogger("azure.monitor.opentelemetry").setLevel(logging.INFO)


# Log startup information
logging.info("=" * 60)
logging.info("BIAL MDA Backend Starting")
logging.info(f"Environment: {'Development' if DEV_MODE else 'Production'}")
logging.info(f"Log Level: {logging.getLevelName(log_level)}")
logging.info("=" * 60)

# --- OBSERVABILITY IMPORTS ---


# Custom span processor to handle export failures gracefully
class ResilientSpanProcessor(BatchSpanProcessor):
    def __init__(self, span_exporter, *args, **kwargs):
        super().__init__(span_exporter, *args, **kwargs)
        self.max_retries = 3
        self.retry_delay = 1.0

    def _export_with_retry(self, spans):
        """Export spans with retry logic and graceful failure handling."""
        for attempt in range(self.max_retries):
            try:
                return self.span_exporter.export(spans)
            except Exception as e:
                if attempt == self.max_retries - 1:
                    # Final attempt failed, log and continue
                    logging.warning(
                        f"Failed to export spans after {self.max_retries} attempts: {e}"
                    )
                    return
                else:
                    # Wait before retry
                    time.sleep(self.retry_delay * (2**attempt))  # Exponential backoff

    def on_end(self, span):
        """Override to use our resilient export method."""
        if span.is_recording():
            self._export_with_retry([span])


# --- Load Environment Variables ---
try:
    load_dotenv()
except Exception as e:
    logging.warning(f"Could not load .env file: {e}")

# --- AZURE MONITOR INITIALIZATION ---
APPLICATIONINSIGHTS_CONNECTION_STRING = os.environ.get(
    "APPLICATIONINSIGHTS_CONNECTION_STRING"
)
LOG_ANALYTICS_WORKSPACE_ID = os.environ.get("LOG_ANALYTICS_WORKSPACE_ID")
LOG_ANALYTICS_WORKSPACE_NAME = os.environ.get("LOG_ANALYTICS_WORKSPACE_NAME")
APPLICATION_INSIGHTS_NAME = os.environ.get(
    "APPLICATION_INSIGHTS_NAME", "bial-mda-platform"
)


def initialize_azure_monitor():
    """Initialize Azure Monitor with OpenTelemetry for comprehensive monitoring."""
    if not APPLICATIONINSIGHTS_CONNECTION_STRING:
        logging.info(
            "Azure Application Insights connection string not configured. Skipping Azure Monitor initialization."
        )
        return False

    try:
        # Configure Azure Monitor with OpenTelemetry
        configure_azure_monitor(
            connection_string=APPLICATIONINSIGHTS_CONNECTION_STRING,
            enable_live_metrics=True,
            enable_standard_metrics=True,
            resource_attributes={
                "service.name": APPLICATION_INSIGHTS_NAME,
                "service.version": "1.0.0",
                "deployment.environment": "development" if DEV_MODE else "production",
            },
        )

        # Instrument other libraries (FastAPI will be instrumented after app creation)
        RequestsInstrumentor().instrument()
        # Enable logging instrumentation so structured logs flow to traces
        LoggingInstrumentor().instrument(set_logging_format=True)
        PymongoInstrumentor().instrument()

        # Instrument OpenAI (use OpenTelemetry version instead of openinference)
        try:
            OpenTelemetryOpenAIInstrumentor().instrument()
        except Exception as e:
            logging.warning(f"Failed to instrument OpenAI with OpenTelemetry: {e}")

        logging.info("Azure Monitor OpenTelemetry initialized successfully")

        # Test if we can create a tracer
        test_tracer = trace.get_tracer(__name__)
        logging.info(f"Test tracer created: {test_tracer}")

        return True

    except Exception as e:
        logging.error(f"Failed to initialize Azure Monitor OpenTelemetry: {e}")
        return False


# Initialize Azure Monitor (without FastAPI instrumentation yet)
azure_monitor_enabled = initialize_azure_monitor()

# --- CONTENT SAFETY CONFIGURATION ---
CONTENT_SAFETY_ENDPOINT = os.environ.get("CONTENT_SAFETY_ENDPOINT")
CONTENT_SAFETY_KEY = os.environ.get("CONTENT_SAFETY_KEY")


class ContentSafetyResult:
    def __init__(self, is_safe: bool, categories: List[str] = None):
        self.is_safe = is_safe
        self.categories = categories or []


async def analyze_content_safety(text: str) -> ContentSafetyResult:
    """Analyze text content for safety violations.
    - Truncates input to service limit (10k chars) to avoid InvalidRequestBody
    - Safely extracts category names regardless of enum/string type
    - MODIFICATION: Made asynchronous by running blocking calls in a thread pool.
    """
    if not CONTENT_SAFETY_ENDPOINT or not CONTENT_SAFETY_KEY:
        logging.warning("Content Safety not configured. Skipping safety check.")
        return ContentSafetyResult(is_safe=True)
    try:
        # Use sync client within thread for async context
        client = ContentSafetyClient(
            endpoint=CONTENT_SAFETY_ENDPOINT,
            credential=AzureKeyCredential(CONTENT_SAFETY_KEY),
        )

        # Service limit: 10_000 characters
        safe_text = text if text is None else text[:10000]
        request = AnalyzeTextOptions(text=safe_text)

        # MODIFICATION: Run blocking SDK call in a thread to avoid blocking the event loop
        response = await asyncio.to_thread(client.analyze_text, request)

        if response.categories_analysis:
            violations = []
            for category in response.categories_analysis:
                if category.severity > 0:
                    try:
                        name = category.category.value
                    except AttributeError:
                        name = str(category.category)
                    violations.append(name)

            if violations:
                logging.warning(f"Content safety violation detected: {violations}")
                return ContentSafetyResult(is_safe=False, categories=violations)

        return ContentSafetyResult(is_safe=True)
    except Exception as e:
        logging.error(f"Content safety analysis failed: {e}")
        # In case of error, allow content through but log the issue
        return ContentSafetyResult(is_safe=True)


# --- LLM OPERATION TRACING DECORATORS ---


def trace_openai_call(client, model_name: str, operation_name: str):
    """
    Wrapper to trace OpenAI API calls with comprehensive metrics.
    MODIFICATION: Now supports async functions.
    """

    def decorator(func):
        @functools.wraps(func)
        async def wrapper(*args, **kwargs):  # MODIFICATION: Changed to async def
            tracer = trace.get_tracer(__name__)
            current_span = trace.get_current_span()

            start_time = time.time()
            operation_id = f"{operation_name}-{int(time.time() * 1000)}"

            # Log structured LLM operation start
            start_log_data = {
                "event": "llm_operation_started",
                "operation_name": operation_name,
                "model_name": model_name,
                "operation_id": operation_id,
                "trace_id": (
                    f"{current_span.get_span_context().trace_id:032x}"
                    if current_span and current_span.get_span_context()
                    else None
                ),
                "user_id": None,
                "user_role": None,
                "query": None,
                "client_type": "openai",
                "usecase": "regulatory",
                "app": "mda",
            }

            try:
                start_log_data["user_id"] = get_value(USER_ID_KEY)
                start_log_data["user_role"] = get_value(USER_ROLE_KEY)
                start_log_data["query"] = get_value(USER_QUERY_KEY)
            except Exception:
                pass

            logging.info(f"LLM Operation Started: {start_log_data}")

            with tracer.start_as_current_span(f"llm_{operation_name}") as llm_span:
                llm_span.set_attribute("llm.operation_name", operation_name)
                llm_span.set_attribute("llm.model_name", model_name)
                llm_span.set_attribute("llm.operation_id", operation_id)
                llm_span.set_attribute("llm.client_type", "openai")

                token_usage = {}
                try:
                    result = await func(*args, **kwargs)  # MODIFICATION: Added await

                    # Process usage data if available
                    if hasattr(result, "usage") and result.usage:
                        token_usage = {
                            "prompt_tokens": result.usage.prompt_tokens or 0,
                            "completion_tokens": result.usage.completion_tokens or 0,
                            "total_tokens": result.usage.total_tokens or 0,
                        }
                        llm_span.set_attribute(
                            "llm.usage.prompt_tokens", token_usage["prompt_tokens"]
                        )
                        llm_span.set_attribute(
                            "llm.usage.completion_tokens",
                            token_usage["completion_tokens"],
                        )
                        llm_span.set_attribute(
                            "llm.usage.total_tokens", token_usage["total_tokens"]
                        )

                    duration = time.time() - start_time
                    throughput = (
                        token_usage.get("total_tokens", 0) / duration
                        if duration > 0
                        else 0
                    )

                    llm_span.set_attribute(
                        "llm.input_tokens", token_usage.get("input_tokens", 0)
                    )
                    llm_span.set_attribute(
                        "llm.output_tokens", token_usage.get("output_tokens", 0)
                    )
                    llm_span.set_attribute(
                        "llm.total_tokens", token_usage.get("total_tokens", 0)
                    )
                    llm_span.set_attribute("llm.latency_seconds", duration)
                    llm_span.set_attribute("llm.latency_ms", round(duration * 1000, 2))
                    llm_span.set_attribute(
                        "llm.throughput_tokens_per_second", round(throughput, 2)
                    )
                    llm_span.set_attribute("llm.success", True)
                    llm_span.set_attribute("llm.status", "completed")

                    user_id = None
                    user_role = None
                    query = None
                    response_preview = None
                    try:
                        user_id = get_value(USER_ID_KEY)
                        user_role = get_value(USER_ROLE_KEY)
                        query = get_value(USER_QUERY_KEY)
                    except Exception:
                        pass

                    try:
                        response_preview = (
                            (result.choices[0].message.content or "")[:200]
                            if hasattr(result, "choices") and result.choices
                            else None
                        )
                    except Exception:
                        response_preview = None

                    completion_log_data = {
                        "event": "llm_operation_completed",
                        "operation_name": operation_name,
                        "model_name": model_name,
                        "operation_id": operation_id,
                        "duration_ms": round(duration * 1000, 2),
                        "input_tokens": token_usage.get("input_tokens", 0),
                        "output_tokens": token_usage.get("output_tokens", 0),
                        "total_tokens": token_usage.get("total_tokens", 0),
                        "throughput_tokens_per_second": round(throughput, 2),
                        "success": True,
                        "error_type": None,
                        "error_message": None,
                        "query": query,
                        "response": response_preview,
                        "user_id": user_id,
                        "user_role": user_role,
                        "client_type": "openai",
                        "usecase": "regulatory",
                        "app": "mda",
                    }
                    logging.info(f"{json.dumps(completion_log_data)}")
                    return result

                except Exception as e:
                    duration = time.time() - start_time

                    llm_span.set_attribute("llm.error_type", type(e).__name__)
                    llm_span.set_attribute("llm.error_message", str(e))
                    llm_span.set_attribute("llm.latency_seconds", duration)
                    llm_span.set_attribute("llm.latency_ms", round(duration * 1000, 2))
                    llm_span.set_attribute("llm.success", False)
                    llm_span.set_attribute("llm.status", "error")

                    llm_span.set_status(Status(StatusCode.ERROR, str(e)))
                    llm_span.record_exception(e)

                    user_id = None
                    user_role = None
                    query = None
                    try:
                        user_id = get_value(USER_ID_KEY)
                        user_role = get_value(USER_ROLE_KEY)
                        query = get_value(USER_QUERY_KEY)
                    except Exception:
                        pass

                    error_log_data = {
                        "event": "llm_operation_error",
                        "operation_name": operation_name,
                        "model_name": model_name,
                        "operation_id": operation_id,
                        "duration_ms": round(duration * 1000, 2),
                        "input_tokens": token_usage.get("input_tokens", 0),
                        "output_tokens": 0,
                        "total_tokens": token_usage.get("input_tokens", 0),
                        "throughput_tokens_per_second": 0,
                        "success": False,
                        "error_type": type(e).__name__,
                        "error_message": str(e),
                        "query": query,
                        "response": None,
                        "user_id": user_id,
                        "user_role": user_role,
                        "client_type": "openai",
                        "usecase": "regulatory",
                        "app": "mda",
                    }
                    logging.error(f"LLM Operation Error: {error_log_data}")

                    raise

        return wrapper

    return decorator


# ==============================================================================
#  NEW: IN-MEMORY MANAGER FOR DEEP RESEARCH & MAPS (NO COSMOS DB REQUIRED)
# ==============================================================================

# ==============================================================================
#  NEW: IN-MEMORY MANAGER FOR DEEP RESEARCH & MAPS (FIXED & COMPLETE)
# ==============================================================================

class InMemoryResearchManager:
    def __init__(self):
        # Stores jobs in RAM: { "user_username": { "job_id": "...", "status": "...", "result": "..." } }
        self.jobs = {} 

    async def run_deep_research_task(self, user_id: str, query: str, context: str):
        """Background task that polls Google and updates the local dictionary."""
        
        if not config.gemini_deepsearch_api_key:
            logging.error("Deep Research API Key missing")
            self._update_job(user_id, "failed", error="API Key missing in .env")
            return

        try:
            client = genai.Client(api_key=config.gemini_deepsearch_api_key)

            # 1. Create Context-Aware Prompt
            full_prompt = (
                f"INTERNAL CONTEXT (From User's Session):\n{context[:20000]}\n\n"
                f"USER QUERY: {query}\n\n"
                "TASK: Verify the internal context against external web sources. "
                "Produce a detailed market report resolving the user's query."
            )
            
            logging.info(f"Starting Google Deep Research for {user_id}...")
            
            # 2. Start the Long-Running Job
            interaction = await asyncio.to_thread(
                client.interactions.create,
                input=full_prompt,
                agent="deep-research-pro-preview-12-2025",
                background=True
            )
            
            # 3. Store Initial State
            self.jobs[user_id] = {
                "job_id": interaction.id,
                "status": "processing",
                "start_time": datetime.now(timezone.utc),
                "query": query
            }

            # 4. Polling Loop (Runs in Background)
            while True:
                await asyncio.sleep(20) # Check every 20 seconds
                
                # --- FIX: Pass ID as positional arg, NOT keyword arg ---
                status_check = await asyncio.to_thread(
                    client.interactions.get, 
                    interaction.id 
                )
                
                if status_check.status == "completed":
                    # JOB DONE: Save result to RAM
                    final_report = status_check.outputs[-1].text
                    self._update_job(user_id, "completed", result=final_report)
                    break
                
                elif status_check.status == "failed":
                    # JOB FAILED
                    err = getattr(status_check, 'error', 'Unknown error')
                    self._update_job(user_id, "failed", error=str(err))
                    break
                    
        except Exception as e:
            logging.error(f"Deep Research Error: {e}", exc_info=True)
            self._update_job(user_id, "failed", error=str(e))

    def _update_job(self, user_id, status, result=None, error=None):
        if user_id in self.jobs:
            self.jobs[user_id]["status"] = status
            if result: self.jobs[user_id]["result"] = result
            if error: self.jobs[user_id]["error"] = error

    def get_job_status(self, user_id):
        return self.jobs.get(user_id)

# Initialize the Global Manager
research_manager = InMemoryResearchManager()

# --- GOOGLE MAPS HELPER FUNCTION ---
def generate_google_map(location_query: str) -> str:
    """Generates a static map image URL."""
    # Use config first, fall back to os.environ
    api_key = getattr(config, 'google_maps_api_key', None) or os.environ.get("GOOGLE_MAPS_API_KEY")
    
    if not api_key:
        return "**[Map Error: GOOGLE_MAPS_API_KEY not configured in .env]**"
    
    clean_loc = urllib.parse.quote(location_query)
    map_url = (
        f"https://maps.googleapis.com/maps/api/staticmap?"
        f"center={clean_loc}&zoom=12&size=600x400&maptype=roadmap"
        f"&markers=color:red%7Clabel:A%7C{clean_loc}"
        f"&key={api_key}"
    )
    return f"\n\n![Map of {location_query}]({map_url})\n\n"

# --- FastAPI App Initialization ---
app = FastAPI(
    title="BIAL Regulatory Assistant API",
    description="A complete, production-grade API for the BIAL multi-agent regulatory analysis platform.",
    version="5.0.0-final",
)


# --- ADD THIS GLOBAL EXCEPTION HANDLER ---
@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    """
    Catches all unhandled exceptions and returns a standardized,
    non-revealing JSON 500 error. This prevents leaking
    internal exception details to the user, fixing the
    "Improper Error Handling" finding.
    """
    # Log the full error to your console/Azure Monitor for debugging
    logging.error(
        f"Global unhandled exception for request {request.url}: {exc}", exc_info=True
    )

    # Return a generic response to the client
    return JSONResponse(
        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        content={
            "error": "InternalServerError",
            "message": "An unexpected internal error occurred. Please try again later.",
            "request_id": str(uuid.uuid4()),  # Provide a unique ID for tracing
        },
    )


# --- END OF GLOBAL EXCEPTION HANDLER ---

# Instrument FastAPI after app creation
if azure_monitor_enabled:
    try:
        FastAPIInstrumentor.instrument_app(app)
        logging.info("FastAPI instrumentation applied successfully")
    except Exception as e:
        logging.warning(f"Failed to instrument FastAPI: {e}")

# Log middleware registration
logging.info("OpenTelemetry request tracking middleware registered")


# --- OPENTELEMETRY REQUEST TRACKING MIDDLEWARE ---
@app.middleware("http")
async def track_requests_with_opentelemetry(request: Request, call_next):
    """Track HTTP requests with OpenTelemetry spans for Azure Monitor integration."""
    tracer = trace.get_tracer(__name__)

    request_id = f"{int(time.time() * 1000)}-{request.method}-{request.url.path.replace('/', '_')}"

    with tracer.start_as_current_span(f"{request.method} {request.url.path}") as span:
        span.set_attribute("http.method", request.method)
        span.set_attribute("http.url", str(request.url))
        span.set_attribute("http.scheme", request.url.scheme)
        span.set_attribute("http.host", request.headers.get("host", "unknown"))
        span.set_attribute("http.target", request.url.path)
        span.set_attribute(
            "http.user_agent", request.headers.get("user-agent", "unknown")
        )
        span.set_attribute("http.request_id", request_id)

        user_info = {}
        try:
            if hasattr(request.state, "current_user") and request.state.current_user:
                user_info = {
                    "user_id": request.state.current_user.username,
                    "user_role": request.state.current_user.role or "user",
                }
                span.set_attribute("enduser.id", user_info["user_id"])
                span.set_attribute("user.role", user_info["user_role"])
                span.set_attribute("chat.user_id", user_info["user_id"])

                attach(set_value(USER_ID_KEY, user_info["user_id"]))
                attach(set_value(USER_ROLE_KEY, user_info["user_role"]))
            else:
                auth_header = request.headers.get("Authorization")
                if auth_header and auth_header.startswith("Bearer "):
                    token = auth_header.split(" ")[1]
                    try:
                        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
                        username = payload.get("sub")
                        if username:
                            db_user = await get_user_from_db(
                                username
                            )  # MODIFICATION: await async db call
                            user_role = (
                                getattr(db_user, "role", "user") if db_user else "user"
                            )

                            user_info = {"user_id": username, "user_role": user_role}
                            span.set_attribute("enduser.id", user_info["user_id"])
                            span.set_attribute("user.role", user_info["user_role"])
                            span.set_attribute("chat.user_id", user_info["user_id"])

                            attach(set_value(USER_ID_KEY, user_info["user_id"]))
                            attach(set_value(USER_ROLE_KEY, user_info["user_role"]))
                    except Exception:
                        pass
        except Exception:
            pass

        start_time = time.time()

        start_log_data = {
            "event": "http_request_started",
            "request_id": request_id,
            "method": request.method,
            "path": request.url.path,
            "trace_id": (
                f"{span.get_span_context().trace_id:032x}"
                if span.get_span_context()
                else None
            ),
            "user_id": user_info.get("user_id"),
            "user_role": user_info.get("user_role"),
            "client_type": "openai",
            "usecase": "regulatory",
            "app": "mda",
        }
        logging.info(f"HTTP Request Started: {start_log_data}")

        try:
            response = await call_next(request)  # Process the request
            duration = time.time() - start_time

            # Set response attributes
            span.set_attribute("http.status_code", response.status_code)
            span.set_attribute("duration.ms", round(duration * 1000, 2))

            # Set status based on HTTP code
            if response.status_code >= 400:
                span.set_status(
                    Status(StatusCode.ERROR, f"HTTP Error {response.status_code}")
                )
            else:
                span.set_status(Status(StatusCode.OK))

            log_data = {
                "event": "http_request_completed",
                "request_id": request_id,
                "method": request.method,
                "path": request.url.path,
                "status_code": response.status_code,
                "duration_ms": round(duration * 1000, 2),
                "success": response.status_code < 400,
                "trace_id": (
                    f"{span.get_span_context().trace_id:032x}"
                    if span.get_span_context()
                    else None
                ),
                "user_id": user_info.get("user_id"),
                "user_role": user_info.get("user_role"),
                "client_type": "openai",
                "usecase": "regulatory",
                "app": "mda",
            }
            logging.info(f"HTTP Request Completed: {log_data}")

            return response
        except Exception as e:
            duration = time.time() - start_time

            span.set_attribute("http.status_code", 500)
            span.set_attribute("http.response.duration", duration)
            span.set_attribute("http.response.success", False)
            span.set_attribute("error.type", type(e).__name__)
            span.set_attribute("error.message", str(e))
            span.set_attribute(
                "error.stack_trace", str(e.__traceback__) if e.__traceback__ else ""
            )

            span.set_status(Status(StatusCode.ERROR, str(e)))
            span.record_exception(e)

            log_data = {
                "event": "http_request_error",
                "request_id": request_id,
                "method": request.method,
                "path": request.url.path,
                "status_code": 500,
                "duration_ms": round(duration * 1000, 2),
                "success": False,
                "error_type": type(e).__name__,
                "error_message": str(e),
                "trace_id": (
                    f"{span.get_span_context().trace_id:032x}"
                    if span.get_span_context()
                    else None
                ),
                "user_id": user_info.get("user_id"),
                "user_role": user_info.get("user_role"),
                "client_type": "openai",
                "usecase": "regulatory",
                "app": "mda",
            }
            logging.error(f"HTTP Request Error: {log_data}")

            raise


# --- CORS Middleware ---
# ALLOW ALL ORIGINS (The Nuclear Fix)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # <--- Change this to allow everything
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- ADD THIS MIDDLEWARE (Replaces any previous version) ---
@app.middleware("http")
async def add_security_headers(request: Request, call_next):
    """
    Middleware to add crucial security headers to all responses
    and set our custom server identity.
    """
    response = await call_next(request)

    # 1. Fix for Missing Security Headers
    response.headers["Content-Security-Policy"] = (
        "default-src 'self'; script-src 'self'; object-src 'none'; frame-ancestors 'self';"
    )
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "SAMEORIGIN"
    response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    response.headers["Strict-Transport-Security"] = (
        "max-age=31536000; includeSubDomains"
    )

    # 2. Remove uvicorn Server header disclosure and set custom server identity
    # Starlette headers are case-insensitive, so we need to remove all variations
    # Remove any existing Server header by iterating through headers
    headers_to_remove = []
    for key in response.headers.keys():
        if key.lower() == "server":
            headers_to_remove.append(key)
    for key in headers_to_remove:
        del response.headers[key]
    # Remove server header to hide backend information
    if "server" in response.headers:
        del response.headers["server"]

    return response


# --- END OF NEW MIDDLEWARE ---


# --- AZURE MONITOR USER TRACKING MIDDLEWARE ---
@app.middleware("http")
async def add_user_to_trace(request: Request, call_next):
    """Middleware to add the current user's ID to telemetry traces for better observability."""
    response = await call_next(request)
    try:
        if hasattr(request.state, "current_user"):
            user = request.state.current_user
            span = trace.get_current_span()
            if span.is_recording():
                span.set_attribute("enduser.id", user.username)
    except Exception as e:
        logging.warning(f"Could not set user in trace: {e}")
    return response


# --- CONFIGURATIONS ---
SECRET_KEY = os.environ.get("SECRET_KEY", "a_very_secret_key_for_production")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60
DEV_MODE = os.environ.get("DEV_MODE", "false").lower() == "true"
TENANT_ID = os.environ.get("TENANT_ID")
CLIENT_ID = os.environ.get("CLIENT_ID")
ISSUER_URL = f"https://login.microsoftonline.com/{TENANT_ID}/v2.0"
JWKS_URL = f"https://login.microsoftonline.com/{TENANT_ID}/discovery/v2.0/keys"
PHOENIX_COLLECTOR_ENDPOINT = os.environ.get("PHOENIX_COLLECTOR_ENDPOINT")
PHOENIX_API_KEY = os.environ.get("PHOENIX_API_KEY")
ENABLE_TELEMETRY = os.environ.get("ENABLE_TELEMETRY", "true").lower() == "true"
ADMIN_API_KEY = os.environ.get("ADMIN_API_KEY")
API_SECRET_KEY = os.environ.get("API_SECRET_KEY")
ALLOW_MAINTENANCE = os.environ.get("ALLOW_MAINTENANCE", "false").lower() == "true"


def initialize_phoenix_tracer():
    """Initializes the Arize Phoenix tracer for LLM observability."""
    if not ENABLE_TELEMETRY:
        logging.info("Telemetry disabled via ENABLE_TELEMETRY environment variable.")
        return

    if not PHOENIX_COLLECTOR_ENDPOINT or not PHOENIX_API_KEY:
        logging.info(
            "Phoenix collector endpoint or API key not configured. Skipping telemetry."
        )
        return

    try:
        resource = Resource(attributes={"service.name": "bial-regulatory-platform"})
        trace_provider = TracerProvider(resource=resource)
        span_exporter = OTLPSpanExporter(
            endpoint=PHOENIX_COLLECTOR_ENDPOINT,
            timeout=3,
            headers={"Authorization": f"Bearer {PHOENIX_API_KEY}"},
        )
        resilient_processor = ResilientSpanProcessor(
            span_exporter,
            max_export_batch_size=256,
            export_timeout_millis=3000,
            schedule_delay_millis=2000,
        )
        trace_provider.add_span_processor(resilient_processor)
        trace.set_tracer_provider(trace_provider)
        OpenAIInstrumentor().instrument()
        logging.info("Arize Phoenix tracer initialized successfully.")

    except Exception as e:
        logging.warning(
            f"Failed to initialize Phoenix tracer: {e}. Continuing without telemetry."
        )


initialize_phoenix_tracer()

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")


def check_creds(cred_value):
    """Helper function to check for missing or placeholder credentials."""
    return not cred_value or "YOUR_" in str(cred_value).upper()


class AzureCredentials(BaseModel):
    """Pydantic model to load and validate all necessary environment variables."""

    search_endpoint: str = os.environ.get("SEARCH_ENDPOINT")
    search_api_key: str = os.environ.get("SEARCH_API_KEY")
    default_search_index_name: str = os.environ.get("DEFAULT_SEARCH_INDEX_NAME")
    default_vector_field_name: str = os.environ.get("DEFAULT_VECTOR_FIELD_NAME")
    default_semantic_config_name: str = os.environ.get("DEFAULT_SEMANTIC_CONFIG_NAME")
    openai_endpoint: str = os.environ.get("OPENAI_ENDPOINT")
    openai_api_version: str = os.environ.get("OPENAI_API_VERSION", "2024-03-01-preview")
    openai_api_key: str = os.environ.get("OPENAI_API_KEY")
    deployment_id: str = os.environ.get("DEPLOYMENT_ID", "gpt-4o-mini")
    mda_deployment_id: str = os.environ.get("AZURE_OPENAI_MDA_DEPLOYMENT_ID", "o3-mini")
    planning_llm_deployment_id: str = os.environ.get(
        "PLANNING_LLM_DEPLOYMENT_ID", "gpt-4o-mini"
    )
    embedding_deployment_id: str = os.environ.get(
        "EMBEDDING_DEPLOYMENT_ID", "text-embedding-3-large"
    )
    cosmos_mongo_connection_string: str = os.environ.get(
        "COSMOS_MONGO_CONNECTION_STRING"
    )
    cosmos_database_name: str = os.environ.get("COSMOS_DATABASE_NAME")
    cosmos_logs_collection: str = os.environ.get("COSMOS_LOGS_COLLECTION")
    cosmos_users_collection: str = os.environ.get("COSMOS_USERS_COLLECTION", "Users")
    cosmos_feedback_collection: str = os.environ.get(
        "COSMOS_FEEDBACK_COLLECTION", "Feedback"
    )
    bing_search_api_key: str = os.environ.get("BING_SEARCH_API_KEY")
    bing_search_endpoint: str = os.environ.get("BING_SEARCH_ENDPOINT")
    serpapi_api_key: str = os.environ.get("SERPAPI_API_KEY")
    content_safety_endpoint: str = os.environ.get("CONTENT_SAFETY_ENDPOINT")
    content_safety_key: str = os.environ.get("CONTENT_SAFETY_KEY")
    bing_custom_config_id: str = os.environ.get("BING_CUSTOM_CONFIG_ID", "reg")
    max_output_tokens: int = int(
        os.environ.get("MAX_OUTPUT_TOKENS", 16000)
    )  # Default to 4096
    # Azure Blob Storage credentials for SAS URL generation
    azure_storage_account_name: str = os.environ.get("AZURE_STORAGE_ACCOUNT_NAME")
    azure_storage_account_key: str = os.environ.get("AZURE_STORAGE_ACCOUNT_KEY")
    gemini_deepsearch_api_key: str = os.environ.get("GEMINI_DEEPSEARCH_API_KEY")
    google_maps_api_key: str = os.environ.get("GOOGLE_MAPS_API_KEY")
    


config = AzureCredentials()


def parse_blob_url(blob_url: str) -> tuple:
    """
    Parses an Azure Blob Storage URL to extract account name, container, and blob name.

    Args:
        blob_url: The full blob storage URL (e.g., https://account.blob.core.windows.net/container/blob.pdf)

    Returns:
        Tuple of (account_name, container_name, blob_name) or (None, None, None) if parsing fails
    """
    try:
        parsed_url = urllib.parse.urlparse(blob_url)

        # Extract account name from hostname (e.g., "account.blob.core.windows.net" -> "account")
        hostname_parts = parsed_url.hostname.split(".")
        if len(hostname_parts) < 1 or "blob" not in parsed_url.hostname:
            logging.warning(f"Invalid blob URL format: {blob_url}")
            return (None, None, None)

        account_name = hostname_parts[0]

        # Extract container and blob path from the URL path
        path_parts = parsed_url.path.strip("/").split("/", 1)
        if len(path_parts) < 2:
            logging.warning(f"Could not parse container/blob from URL: {blob_url}")
            return (None, None, None)

        container_name = path_parts[0]
        blob_name = path_parts[1]

        return (account_name, container_name, blob_name)
    except Exception as e:
        logging.error(f"Error parsing blob URL {blob_url}: {e}", exc_info=True)
        return (None, None, None)


def generate_blob_proxy_url(blob_url: str) -> str:
    """
    Generates a proxy URL for an Azure Blob Storage URL.
    This URL will be used by the frontend to access blobs through the backend proxy.

    Args:
        blob_url: The full blob storage URL (e.g., https://account.blob.core.windows.net/container/blob.pdf)

    Returns:
        The proxy URL that the frontend can use to access the blob
    """
    try:
        # Parse the blob URL to extract container and blob name
        account_name, container_name, blob_name = parse_blob_url(blob_url)

        if not account_name or not container_name or not blob_name:
            logging.warning(
                f"Could not generate proxy URL for {blob_url}, returning original"
            )
            return blob_url

        # Encode the blob URL components for the proxy endpoint
        # Use base64 encoding to safely pass the blob identifier
        blob_identifier = (
            base64.urlsafe_b64encode(f"{container_name}/{blob_name}".encode("utf-8"))
            .decode("utf-8")
            .rstrip("=")
        )

        # Generate proxy URL (will be handled by the /blob-proxy endpoint)
        proxy_url = f"/blob-proxy?blob={blob_identifier}"
        logging.debug(f"Generated proxy URL for blob: {blob_name[:50]}...")
        return proxy_url

    except Exception as e:
        logging.error(f"Error generating proxy URL for {blob_url}: {e}", exc_info=True)
        return blob_url


class ServiceClients:
    def __init__(self):
        """Initializes all external service clients."""
        global config  # Ensure access to the global config object

        # OpenAI Clients (Sync used with asyncio.to_thread)
        self.synthesis_openai_client = None  # Initialize as None
        self.planning_openai_client = None  # Initialize as None
        self.search_query_embeddings_model = None  # Initialize as None
        # Check credentials BEFORE trying to initialize
        if not check_creds(config.openai_api_key) and not check_creds(
            config.openai_endpoint
        ):
            try:
                # Attempt to initialize OpenAI clients
                self.synthesis_openai_client = AzureOpenAI(
                    api_key=config.openai_api_key,
                    azure_endpoint=config.openai_endpoint,
                    api_version=config.openai_api_version,
                )
                # Assume planning client uses the same instance unless configured differently
                self.planning_openai_client = self.synthesis_openai_client
                self.search_query_embeddings_model = AzureOpenAIEmbeddings(
                    azure_deployment=config.embedding_deployment_id,
                    azure_endpoint=config.openai_endpoint,
                    api_key=config.openai_api_key,
                    api_version=config.openai_api_version,
                )
                logging.info("OpenAI clients initialized successfully.")
            except Exception as e_openai:
                # Log the specific error if initialization fails
                logging.error(
                    f"Failed to initialize OpenAI clients: {e_openai}", exc_info=True
                )
                # Ensure attributes remain None on failure
                self.synthesis_openai_client = None
                self.planning_openai_client = None
                self.search_query_embeddings_model = None
        else:
            # Log if credentials check failed
            logging.warning(
                "OpenAI credentials missing or invalid (check_creds failed), clients not initialized."
            )
            # Ensure attributes are None
            self.synthesis_openai_client = None
            self.planning_openai_client = None
            self.search_query_embeddings_model = None

        # Sync Search Client Initialization
        self.sync_search_client = None  # Initialize as None first
        if not check_creds(config.search_endpoint) and not check_creds(
            config.search_api_key
        ):
            try:
                self.sync_search_client = SyncSearchClient(
                    endpoint=config.search_endpoint,
                    index_name=config.default_search_index_name,
                    credential=AzureKeyCredential(config.search_api_key),
                )
                logging.info(
                    f"Sync Azure Search client initialized for index '{config.default_search_index_name}'."
                )
            except Exception as e_sync_search:
                logging.error(
                    f"Failed to initialize Sync Azure Search client: {e_sync_search}",
                    exc_info=True,
                )
                self.sync_search_client = None  # It remains None after failure
        else:
            logging.warning(
                "Azure Search credentials missing, Sync client not initialized."
            )

        # MongoDB Client (Sync)
        self.mongo_client = None
        if not check_creds(config.cosmos_mongo_connection_string):
            try:
                self.mongo_client = MongoClient(
                    config.cosmos_mongo_connection_string, serverSelectionTimeoutMS=5000
                )
                self.mongo_client.admin.command("ping")  # Test connection
                logging.info("MongoDB client initialized and connection verified.")
            except Exception as e_mongo:
                logging.error(
                    f"Failed to initialize MongoDB client: {e_mongo}", exc_info=True
                )
                self.mongo_client = None
        else:
            logging.warning(
                "MongoDB connection string missing, client not initialized."
            )

        # Content Safety Client (Sync)
        self.content_safety_client = None
        if not check_creds(config.content_safety_endpoint) and not check_creds(
            config.content_safety_key
        ):
            try:
                self.content_safety_client = ContentSafetyClient(
                    config.content_safety_endpoint,
                    AzureKeyCredential(config.content_safety_key),
                )
                logging.info("Content Safety client initialized.")
            except Exception as e_cs:
                logging.error(
                    f"Failed to initialize Content Safety client: {e_cs}", exc_info=True
                )
                self.content_safety_client = None
        else:
            logging.warning(
                "Content Safety credentials missing, client not initialized."
            )

        # Async HTTP Client (for web search)
        self.async_http_client = httpx.AsyncClient(timeout=30.0)

        # JWKS Client (Sync)
        self.jwks_client = None
        if JWKS_URL:  # Check if JWKS_URL is defined and not None/empty
            try:
                # Consider adding caching/lifespan if appropriate for pyjwt
                self.jwks_client = pyjwt.PyJWKClient(
                    JWKS_URL
                )  # Removed cache_keys/lifespan for simplicity if causing issues
                logging.info("JWKS client initialized.")
            except Exception as e_jwks:
                logging.error(
                    f"Could not fetch/initialize JWKS client from Microsoft: {e_jwks}",
                    exc_info=True,
                )
        else:
            logging.warning("JWKS URL not configured, SSO validation might fail.")


# <<< The ServiceClients class definition ends here

# --- !!! ADD THIS LINE AT GLOBAL SCOPE !!! ---
clients = ServiceClients()
# --- !!! END ADDED LINE !!! ---


# --- Pydantic Models for API Requests and Responses ---
class ChatMessage(BaseModel):
    role: str
    content: str


class ConversationalChatResponse(BaseModel):
    answer: str
    plan: List[str]
    sources: List[Dict[str, Any]]
    source: str  # Added in original


class ChatRequest(BaseModel):  # Kept from original
    question: str
    history: List[ChatMessage] = []


class User(BaseModel):  # Kept from original
    username: str
    role: Optional[str] = None
    isActive: Optional[bool] = True
    min_token_iat: Optional[datetime] = None


class UserInDB(User):  # Kept from original
    hashed_password: str


class Token(BaseModel):  # Kept from original
    access_token: str
    token_type: str


class TokenData(BaseModel):  # Kept from original
    username: Optional[str] = None


class UserCreate(BaseModel):  # Kept from original
    id: Optional[str] = None
    username: str
    password: str
    role: Optional[str] = "user"
    isActive: Optional[bool] = True
    lastLogin: Optional[datetime] = None
    createdAt: Optional[datetime] = None


class AdminCreateUserRequest(BaseModel):  # Kept from original
    id: Optional[str] = None
    username: str
    password: Optional[str] = None
    password_hash: Optional[str] = None
    role: Optional[str] = "user"
    isActive: Optional[bool] = True
    createdAt: Optional[datetime] = None
    lastLogin: Optional[datetime] = None


class MaintenanceSetPasswordRequest(BaseModel):  # Kept from original
    username: str
    new_password: str


class FeedbackRequest(BaseModel):  # Kept from original
    question: str
    answer: str
    feedback: str


class DownloadRequest(BaseModel):
    html_content: str


class SSOLoginRequest(BaseModel):
    sso_token: str


class RefineReportRequest(BaseModel):  # Kept from original
    original_report: str
    new_info: str


class RefineReportResponse(BaseModel):  # Kept from original
    refined_report: str


class CSVAnalysisResponse(BaseModel):
    session_id: str
    text_output: Optional[str] = None
    image_output: Optional[str] = None
    error: Optional[str] = None


# --- NEW USER MANAGEMENT MODELS ---
class NewUserCreate(BaseModel):
    username: str
    email: str
    password: str
    first_name: str
    last_name: str
    role: Optional[str] = "user"


class NewUserResponse(BaseModel):
    success: bool
    message: str
    username: str
    action: str
    details: Dict[str, Any]


class NewUserUpdate(BaseModel):
    username: Optional[str] = None
    email: Optional[str] = None
    password: Optional[str] = None
    first_name: Optional[str] = None
    last_name: Optional[str] = None
    role: Optional[str] = None
    permissions: Optional[Dict[str, Any]] = None


class NewUserDisableRequest(BaseModel):
    username: str
    disable: bool = True


class NewUserListResponse(BaseModel):
    id: str
    username: str
    email: str
    first_name: str
    last_name: str
    role: str
    permissions: Dict[str, Any]
    account_locked: bool
    created_at: Optional[datetime] = None
    last_login_at: Optional[datetime] = None


class RoleDetails(BaseModel):  # Kept from original
    role_name: str
    permissions: List[str]
    description: str


csv_sessions = {}

# --- FULL, UNABRIDGED SYSTEM PROMPT ---
# SHARED_SYSTEM_PROMPT = """You are an expert AI assistant for analyzing Multi-Year Tariff (MYT) submissions for the Airports Economic Regulatory Authority (AERA). Your primary function is to provide detailed, accurate, and context-aware analysis based on regulatory documents. Your final response should be comprehensive and well-structured.

# **Core Directives:**
# 1.  **Strictly Adhere to Context:** Base your answers *only* on the information provided in the aggregated context from search results. Do not use external knowledge unless a web search was explicitly performed.
# 2.  **Precise Terminology:** When discussing the Authority's stance, use the exact terminology found in the source documents (e.g., "approved," "decided," "proposed," "considered"). If the user's term differs from the document's term, you must highlight this discrepancy. For example: "You asked for 'approved' figures. The document describes what was 'decided by the authority' as follows..."
# 3.  **Mandatory Table Referencing:** For any data, figures, or claims extracted from a table, you *must* cite the corresponding table number in your response (e.g., "The Authority approved Aeronautical Revenue of ₹1,500 Cr for FY 2024-25 (Table 15).").
# 4.  **Synthesize, Don't Just List:** Combine information from different search results to form a cohesive, narrative answer. Do not simply list out chunks of text.
# 5.  **Professional Formatting:** Use HTML for tables (`<table border='1'>`) to present numerical data clearly. Use headings and bullet points to structure your response.
# 6.  **Cite Sources:** At the end of your answer, include a 'References:' section listing the source documents from which the information was derived.
# 7.  If the user askes for data in any programming in code or file format (eg. .py, .jsx, .html, etc), do not provide the code or file strictly.
# 8.  If anybody is asking about kyndi and sinequa, just ignore answering it strictly.
# """

SHARED_SYSTEM_PROMPT = """You are an expert AI assistant for conducting market analysis and comparing revenues across multiple sectors and companies. Your primary function is to provide detailed, accurate, and context-aware analysis based on market data, financial reports, and business metrics. Your final response should be comprehensive, well-structured, and insightful.

**Core Directives:**
1.  **Strictly Adhere to Context:** Base your answers *only* on the information provided in the aggregated data and reports. Do not use external knowledge unless a web search was explicitly performed.
2.  **Precise Terminology:** When discussing market data or financial performance, use the exact terminology found in the source documents (e.g., "revenue growth," "market share," "projected," "quarterly results"). If the user's term differs from the document's term, you must highlight this discrepancy. For example: "You asked for 'net income.' The document refers to 'adjusted EBITDA' as follows..."
3.  **Mandatory Table Referencing:** For any data, figures, or claims extracted from a table, you *must* cite the corresponding table number in your response (e.g., "The company reported a 5% increase in revenue for Q3 2024 (Table 8).").
4.  **Synthesize, Don't Just List:** Combine information from different reports or data points to form a cohesive, insightful analysis. Do not simply list out raw data.
5.  **Professional Formatting:** Use HTML for tables (`<table border='1'>`) to present numerical data clearly. Use headings and bullet points to structure your response.
6.  **Revenue and Market Metrics Comparison:** Where appropriate, compare key financial and market metrics (e.g., revenue growth, profitability, market share) across different sectors or companies to provide a comparative analysis.
7.  **Cite Sources:** At the end of your answer, include a 'References:' section listing the source documents from which the information was derived.
8.  **Do not provide programming code or files** in formats like .py, .jsx, .html, etc.
9.  **Ignore requests regarding Kyndi and Sinequa.**
"""


# --- AUTHENTICATION & USER MANAGEMENT ---
# ... (Functions kept: verify_password, get_password_hash, create_access_token) ...
def verify_password(plain, hashed):
    return pwd_context.verify(plain, hashed)


def get_password_hash(password):
    return pwd_context.hash(password)


# In mda.py (around line 1032) AND ff.py (around line 957)


def create_access_token(
    data: dict,
    expires_delta: Optional[timedelta] = None,
    iat_time: Optional[datetime] = None,
):
    to_encode = data.copy()

    # Use the provided time, or get a new one if not provided
    now = iat_time if iat_time else datetime.now(timezone.utc)

    expire = now + (expires_delta or timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES))
    to_encode.update(
        {
            "exp": expire,
            "jti": str(uuid.uuid4()),
            "iat": now,
        }
    )
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)


async def get_user_from_db(username: str) -> Optional[UserInDB]:
    """Fetches user from DB asynchronously."""
    global clients  # Explicitly use the global clients object
    if not clients.mongo_client:
        logging.warning(
            f"Attempted DB lookup for {username}, but DB client is not configured."
        )
        return None

    def db_call_sync():
        try:
            user_doc = clients.mongo_client[config.cosmos_database_name][
                config.cosmos_users_collection
            ].find_one({"username": username})
            if user_doc:
                user_doc.setdefault(
                    "hashed_password", user_doc.get("password_hash")
                )  # Compatibility
                user_doc.setdefault("role", "user")
                user_doc.setdefault("isActive", True)
                if "_id" in user_doc and "id" not in user_doc:
                    user_doc["id"] = str(user_doc["_id"])
                # Ensure all required fields for UserInDB are present or provide defaults
                user_doc.setdefault("username", username)  # Ensure username is present
                if "hashed_password" not in user_doc:
                    logging.error(
                        f"User {username} found in DB but missing hashed_password."
                    )
                    return None  # Cannot authenticate without password hash
                return UserInDB(**user_doc)
            return None
        except Exception as e_db_lookup:
            logging.error(
                f"Error during DB lookup for {username}: {e_db_lookup}", exc_info=True
            )
            return None

    return await asyncio.to_thread(db_call_sync)


async def get_current_user_prod(token: str = Depends(oauth2_scheme)) -> User:
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        jti: str = payload.get("jti")
        iat_timestamp: int = payload.get("iat")

        if username is None or jti is None or iat_timestamp is None:
            logging.warning("Token missing 'sub', 'jti', or 'iat' claim.")
            raise credentials_exception

        # Convert token's 'iat' from timestamp to datetime (OFFSET-AWARE)
        token_iat_time = datetime.fromtimestamp(iat_timestamp, timezone.utc)

        # --- BLOCKLIST CHECK (Scenario 1) ---
        if clients.mongo_client:
            blocklist_collection = clients.mongo_client[config.cosmos_database_name][
                "token_blocklist"
            ]

            # --- START OF THE FIX ---
            # This check reliably determines if the client is ASYNC (ff.py)
            # or SYNC (mda.py) by checking its class name.
            if "AsyncIOMotorClient" in str(type(clients.mongo_client)):
                # This is ff.py (async)
                is_blocklisted = await blocklist_collection.find_one({"jti": jti})
            else:
                # This is mda.py (sync)
                is_blocklisted = await asyncio.to_thread(
                    blocklist_collection.find_one, {"jti": jti}
                )
            # --- END OF THE FIX ---

            if is_blocklisted:
                logging.warning(f"Revoked token (jti: {jti}) was used by {username}.")
                raise HTTPException(
                    status_code=401, detail="Token has been revoked (logged out)."
                )
        # --- END BLOCKLIST CHECK ---

    except JWTError as e:
        logging.warning(f"JWTError: {e}")
        raise credentials_exception

    user = await get_user_from_db(username)
    if user is None:
        raise credentials_exception

    # --- CONCURRENCY CHECK (Scenario 2) ---
    if user.min_token_iat:
        user_min_token_aware = user.min_token_iat.replace(tzinfo=timezone.utc)
        if token_iat_time < user_min_token_aware:
            logging.warning(
                f"Old token used by {username}. Token iat: {token_iat_time}, Required iat: {user_min_token_aware}"
            )
            raise HTTPException(
                status_code=401,
                detail="This token is no longer valid. Please log in again.",
            )
    # --- END CONCURRENCY CHECK ---

    # --- 'IS ACTIVE' CHECK ---
    if hasattr(user, "isActive") and not user.isActive:
        logging.warning(f"Attempted login by inactive user: {username}")
        raise HTTPException(status_code=401, detail="User account is disabled.")
    # --- END 'IS ACTIVE' CHECK ---

    return User(
        username=user.username,
        role=user.role,
        isActive=user.isActive,
        min_token_iat=user.min_token_iat,
    )


async def get_current_user_dev_and_set_state(request: Request):
    user = User(username="dev_user")
    request.state.current_user = user
    return user


async def get_current_user_and_set_state(
    request: Request, token: str = Depends(oauth2_scheme)
):
    user = await get_current_user_prod(token)
    request.state.current_user = user
    return user


auth_dependency = (
    get_current_user_dev_and_set_state if DEV_MODE else get_current_user_and_set_state
)


def require_api_secret(request: Request):
    if not API_SECRET_KEY:
        logging.error("API_SECRET_KEY environment variable is not set")
        raise HTTPException(status_code=503, detail="API Secret not configured")

    provided = request.headers.get("x-api-secret") or request.headers.get(
        "X-API-Secret"
    )
    if not provided:
        logging.warning("X-API-Secret header not provided in request")
        raise HTTPException(status_code=401, detail="X-API-Secret header required")

    if provided != API_SECRET_KEY:
        logging.warning(
            f"Invalid API secret provided. Expected: {API_SECRET_KEY[:8]}..., Got: {provided[:8] if provided else 'None'}..."
        )
        raise HTTPException(status_code=401, detail="Unauthorized - Invalid API Secret")


# --- LOGGING FUNCTION ---
async def log_interaction(user: str, question: str, answer: str, duration: float):
    """MODIFICATION: Made asynchronous to avoid blocking."""
    if not clients.mongo_client:
        logging.warning(
            "Could not log interaction: Database service is not configured."
        )
        return

    def db_call():
        try:
            db = clients.mongo_client[config.cosmos_database_name]
            log_collection = db[config.cosmos_logs_collection]
            current_time = datetime.now(timezone.utc)
            log_document = {
                "timestamp": current_time,
                "logDate": current_time,
                "user": user,
                "question": question,
                "answer": answer[:5000] if answer else "",
                "answer_length": len(answer) if answer else 0,
                "duration": duration,
                "processed": False,
            }
            logging.info(
                f"Attempting to insert into {config.cosmos_logs_collection} for user: {user}"
            )
            result = log_collection.insert_one(log_document)
            logging.info(
                f"Successfully logged interaction with ID: {result.inserted_id}"
            )
            return True
        except Exception as e:
            logging.error(
                f"FAILED to write to {config.cosmos_logs_collection}: {e}",
                exc_info=True,
            )
            try:
                if clients.mongo_client:
                    db = clients.mongo_client[config.cosmos_database_name]
                    collections = db.list_collection_names()
                    logging.error(f"Available collections: {collections}")
                    logging.error(
                        f"Target collection '{config.cosmos_logs_collection}' exists: {config.cosmos_logs_collection in collections}"
                    )
            except Exception as debug_error:
                logging.error(f"Debug info failed: {debug_error}")
            return False

    await asyncio.to_thread(db_call)


# --- DOCUMENT EXTRACTION (Enhanced version w/ Tables) ---
async def extract_text_from_docx(file: UploadFile) -> str:
    """Extracts text from DOCX including tables (async)."""
    try:
        content = await file.read()

        def process_docx_sync():
            document = docx.Document(io.BytesIO(content))
            full_text = []
            for element in document.element.body:
                if element.tag.endswith("p"):
                    para = next(
                        (p for p in document.paragraphs if p._element is element), None
                    )
                    if para:
                        full_text.append(para.text.strip())  # Strip whitespace
                elif element.tag.endswith("tbl"):
                    try:
                        table = next(
                            (tbl for tbl in document.tables if tbl._element is element),
                            None,
                        )
                        if table:
                            # Extract table data row by row, joining cells with '|'
                            table_rows = []
                            for row in table.rows:
                                row_text = " | ".join(
                                    [
                                        cell.text.strip().replace("\n", " ")
                                        for cell in row.cells
                                    ]
                                )
                                table_rows.append(row_text)
                            # Add header separator for clarity if table is not empty
                            if table_rows:
                                header = table_rows[0]
                                separator = " | ".join(
                                    ["---"] * len(header.split(" | "))
                                )
                                # Prepend "[TABLE START]" and append "[TABLE END]" markers
                                full_text.append(
                                    "[TABLE START]\n"
                                    + "\n".join([header, separator] + table_rows[1:])
                                    + "\n[TABLE END]"
                                )
                            # else: # Handle empty table if needed
                            #     full_text.append("[Empty Table]")
                    except Exception as e_table:
                        logging.error(
                            f"Error parsing a table in DOCX: {e_table}", exc_info=True
                        )
                        full_text.append("[Error extracting table content]")
            # Join all extracted parts, filtering empty strings
            return "\n\n".join(filter(None, full_text))

        # Run sync processing in thread
        return await asyncio.to_thread(process_docx_sync)
    except Exception as e:
        logging.error(
            f"Error reading docx file {getattr(file, 'filename', 'unknown')}: {e}",
            exc_info=True,
        )
        # Raise HTTPException for client feedback
        raise HTTPException(status_code=500, detail=f"Error reading docx file: {e}")


# --- ASYNC Vector Generation (NEW) ---
async def get_query_vector(text: str) -> Optional[List[float]]:
    """Generates a vector embedding asynchronously."""
    if not clients.search_query_embeddings_model:
        logging.warning("Embeddings model not initialized.")
        return None
    try:
        # Run the sync Langchain call in a thread
        return await asyncio.to_thread(
            clients.search_query_embeddings_model.embed_query, text
        )
    except Exception as e:
        logging.error(f"Vector generation error: {e}", exc_info=True)
        return None  # Return None on error


# --- ASYNC Azure Search Query (Enhanced w/ Detailed Refs) ---
async def query_azure_search_async(
    query_text: str, index_name: str, k: int = 5, use_hybrid_semantic: bool = True
):
    """Performs Azure Search query asynchronously and extracts detailed references."""
    context, references_data = "", []
    search_endpoint = config.search_endpoint
    search_api_key = config.search_api_key
    vector_field_name = config.default_vector_field_name
    semantic_config_name = config.default_semantic_config_name

    if check_creds(search_endpoint) or check_creds(search_api_key):
        logging.warning("Azure Search client/credentials not configured.")
        return "Error: Azure Search service is not configured.", []

    # Use the async client directly
    try:
        async_search_client = AsyncSearchClient(
            search_endpoint, index_name, AzureKeyCredential(search_api_key)
        )
        async with async_search_client:  # Use context manager for proper cleanup
            search_kwargs = {
                "search_text": query_text if query_text and query_text.strip() else "*",
                "top": k,
                "include_total_count": True,
                # Explicitly list fields needed for context and referencing
                "select": "content, filepath, url, title, source_reference, document_title, page_number, table_references, figure_references, section_references, section_info, table_info, figure_info",
            }

            # Add hybrid/semantic parameters
            if use_hybrid_semantic:
                query_vector = await get_query_vector(query_text)  # Await async call
                if query_vector and not check_creds(vector_field_name):
                    search_kwargs["vector_queries"] = [
                        VectorizedQuery(
                            vector=query_vector,
                            k_nearest_neighbors=k,
                            fields=vector_field_name,
                        )
                    ]
                # Only set semantic if config name is valid
                if not check_creds(semantic_config_name):
                    search_kwargs.update(
                        {
                            "query_type": "semantic",
                            "semantic_configuration_name": semantic_config_name,
                            "query_caption": "extractive",  # Keep as extractive
                            "query_answer": "extractive",  # Keep as extractive
                        }
                    )
                # If only vector search is possible (no semantic config), don't set query_type='semantic'
                elif query_vector:
                    # If vector exists but no semantic config, default might be simple or vector search needs explicit type?
                    # Check Azure Search SDK docs - often 'simple' is default and vector works with it.
                    search_kwargs.setdefault(
                        "query_type", "simple"
                    )  # Ensure query_type is set if vector is used alone

            # Perform the async search
            logging.debug(f"Executing Azure Search async with params: {search_kwargs}")
            results = await async_search_client.search(**search_kwargs)
            count = await results.get_count()
            logging.info(
                f"Azure Search returned {count} results async for query: '{query_text[:50]}...'"
            )
            if count == 0:
                return "", []

            processed_references = {}
            async for doc in results:
                doc_content = doc.get("content", "")
                context += doc_content + "\n\n"  # Append content with spacing

                # --- Extract Detailed Reference Info ---
                title = doc.get("title", "")
                filepath = doc.get("filepath", "")
                url = doc.get("url", "")
                document_title = doc.get("document_title", "")
                page_number = doc.get("page_number", "")
                table_refs_raw = doc.get("table_references", [])
                figure_refs_raw = doc.get("figure_references", [])
                section_refs_raw = doc.get("section_references", [])

                # Process filepath to generate URL if needed (similar to ff_final.py)
                # All filepaths from search index are https:// blob URLs
                if filepath:
                    # Check if filepath is an Azure Blob Storage URL
                    if (
                        filepath.startswith("https://")
                        and "blob.core.windows.net" in filepath
                    ):
                        # Generate proxy URL for blob storage (secure access without exposing blob URLs)
                        url = generate_blob_proxy_url(filepath)
                    elif filepath.startswith("https://"):
                        # For other HTTPS URLs, use as-is
                        url = filepath
                    else:
                        # If filepath is not https://, log warning and use placeholder
                        # This should not happen if all filepaths are https:// blob URLs
                        logging.warning(
                            f"Unexpected filepath format (not https://): {filepath}. Using placeholder URL."
                        )
                        filename = (
                            document_title
                            or title
                            or f"Result {len(processed_references) + 1}"
                        )
                        url = f"#document-{filename.replace(' ', '-').lower()}"
                elif not url:
                    # If no filepath and no url, create a document anchor
                    filename = (
                        document_title
                        or title
                        or f"Result {len(processed_references) + 1}"
                    )
                    url = f"#document-{filename.replace(' ', '-').lower()}"

                def safe_json_loads(field_name, default_value):
                    raw_value = doc.get(field_name)
                    if isinstance(raw_value, str):
                        try:
                            return json.loads(raw_value)
                        except json.JSONDecodeError:
                            logging.warning(f"Failed to decode JSON for {field_name}")
                            return default_value
                    # Handle cases where it might already be parsed (or None)
                    return raw_value if raw_value is not None else default_value

                section_info = safe_json_loads("section_info", [])
                table_info = safe_json_loads("table_info", [])
                figure_info = safe_json_loads("figure_info", [])

                # Clean up reference lists
                table_references = (
                    [str(t) for t in table_refs_raw if t]
                    if isinstance(table_refs_raw, list)
                    else []
                )
                figure_references = (
                    [str(f) for f in figure_refs_raw if f]
                    if isinstance(figure_refs_raw, list)
                    else []
                )
                section_references = (
                    [str(s) for s in section_refs_raw if s]
                    if isinstance(section_refs_raw, list)
                    else []
                )

                # Construct display name and reference string
                display_name = (
                    document_title
                    or title
                    or (
                        os.path.basename(filepath)
                        if filepath
                        else f"Result {len(processed_references) + 1}"
                    )
                )
                reference_parts = []
                if page_number:
                    reference_parts.append(f"Page {page_number}")
                if section_references:
                    reference_parts.append(f"Sec: {', '.join(section_references)}")
                if table_references:
                    reference_parts.append(f"Table: {', '.join(table_references)}")
                if figure_references:
                    reference_parts.append(f"Fig: {', '.join(figure_references)}")
                detailed_reference_string = (
                    f"{display_name} ({', '.join(reference_parts)})"
                    if reference_parts
                    else display_name
                )

                # Determine a unique key for deduplication
                ref_key = url or detailed_reference_string

                # Store unique references
            if ref_key not in processed_references:
                processed_references[ref_key] = {
                    "filename_or_title": detailed_reference_string,
                    "url": url,
                    # Removed "filepath" from response to hide internal blob storage paths
                    "source_type": "document_chunk",
                    "page_number": page_number,
                    "document_title": document_title,
                    "section_references": section_references,
                    "table_references": table_references,
                    "figure_references": figure_references,
                    "section_info": section_info,
                    "table_info": table_info,
                    "figure_info": figure_info,
                    "content": doc_content,
                    "score": doc.get("@search.score"),
                    "reranker_score": doc.get("@search.reranker_score"),
                }
            references_data = list(processed_references.values())

    except Exception as e:
        logging.error(
            f"Error during Azure Search async for index '{index_name}': {e}",
            exc_info=True,
        )
        return f"Error accessing search index '{index_name}': {e}", []

    return context.strip(), references_data


# --- ORIGINAL SYNC Azure Search (For original chat logic - WITH DETAILED REFS) ---
def query_azure_search_sync(
    query_text: str,
    index_name: str,
    k: int = 5,
    use_hybrid_semantic_search: bool = True,
):
    """Synchronous version for the original chat endpoint.
    MODIFIED to extract and format detailed references."""
    context, references_data = "", []
    search_endpoint = config.search_endpoint
    search_api_key = config.search_api_key
    vector_field_name = config.default_vector_field_name
    semantic_config_name = config.default_semantic_config_name

    if check_creds(search_endpoint) or check_creds(search_api_key):
        logging.warning("Sync Search: Credentials missing.")
        return "Error: Azure Search service is not configured.", []
    if not clients.sync_search_client:  # Check if sync client is initialized
        logging.error("Sync Azure Search client is not available.")
        # Attempt to initialize it here? Or rely on initial setup.
        # For safety, return error if not initialized during startup
        return "Error: Sync Search client not initialized.", []

    # Ensure correct index - Create a NEW client instance for the specific index if needed
    # This avoids issues if multiple indices are used across different calls
    try:
        current_search_client = SyncSearchClient(
            search_endpoint, index_name, AzureKeyCredential(search_api_key)
        )
        logging.debug(f"Using Sync Search client for index '{index_name}'.")
    except Exception as client_err:
        logging.error(
            f"Failed to initialize sync search client for '{index_name}': {client_err}"
        )
        return (
            f"Error initializing sync search client for index '{index_name}': {client_err}",
            [],
        )

    try:
        search_kwargs = {
            "search_text": query_text if query_text and query_text.strip() else "*",
            "top": k,
            "include_total_count": True,
            # Explicitly select fields needed for detailed referencing
            "select": "content, filepath, url, title, source_reference, document_title, page_number, table_references, figure_references, section_references, section_info, table_info, figure_info",
        }
        if use_hybrid_semantic_search:
            # Sync embedding generation
            query_vector = None
            if clients.search_query_embeddings_model:
                try:
                    query_vector = clients.search_query_embeddings_model.embed_query(
                        query_text
                    )
                except Exception as embed_err:
                    logging.warning(
                        f"Sync Search: Failed to generate query vector: {embed_err}"
                    )

            if query_vector and not check_creds(vector_field_name):
                search_kwargs["vector_queries"] = [
                    VectorizedQuery(
                        vector=query_vector,
                        k_nearest_neighbors=k,
                        fields=vector_field_name,
                    )
                ]
            # Only add semantic config if name is valid
            if not check_creds(semantic_config_name):
                search_kwargs.update(
                    {
                        "query_type": "semantic",
                        "semantic_configuration_name": semantic_config_name,
                        "query_caption": "extractive",
                        "query_answer": "extractive",
                    }
                )
            elif (
                query_vector
            ):  # If vector added but no semantic config, ensure type allows vector
                search_kwargs.setdefault(
                    "query_type", "simple"
                )  # Default usually works

        results = current_search_client.search(**search_kwargs)
        count = results.get_count()
        logging.info(
            f"Sync Azure Search returned {count} results for query: '{query_text[:50]}...'"
        )
        if count == 0:
            return "", []

        processed_references = {}
        for doc in results:
            doc_content = doc.get("content", "")
            context += doc_content + "\n\n"

            # --- Extract Detailed Reference Info (same as async version) ---
            title = doc.get("title", "")
            filepath = doc.get("filepath", "")
            url = doc.get("url", "")
            document_title = doc.get("document_title", "")
            page_number = doc.get("page_number", "")
            table_refs_raw = doc.get("table_references", [])
            figure_refs_raw = doc.get("figure_references", [])
            section_refs_raw = doc.get("section_references", [])

            # Process filepath to generate URL if needed (similar to ff_final.py)
            # All filepaths from search index are https:// blob URLs
            if filepath:
                # Check if filepath is an Azure Blob Storage URL
                if (
                    filepath.startswith("https://")
                    and "blob.core.windows.net" in filepath
                ):
                    # Generate proxy URL for blob storage (secure access without exposing blob URLs)
                    url = generate_blob_proxy_url(filepath)
                elif filepath.startswith("https://"):
                    # For other HTTPS URLs, use as-is
                    url = filepath
                else:
                    # If filepath is not https://, log warning and use placeholder
                    # This should not happen if all filepaths are https:// blob URLs
                    logging.warning(
                        f"Unexpected filepath format (not https://): {filepath}. Using placeholder URL."
                    )
                    filename = (
                        document_title
                        or title
                        or f"Result {len(processed_references) + 1}"
                    )
                    url = f"#document-{filename.replace(' ', '-').lower()}"
            elif not url:
                # If no filepath and no url, create a document anchor
                filename = (
                    document_title or title or f"Result {len(processed_references) + 1}"
                )
                url = f"#document-{filename.replace(' ', '-').lower()}"

            def safe_json_loads(field_name, default_value):
                raw_value = doc.get(field_name)
                if isinstance(raw_value, str):
                    try:
                        return json.loads(raw_value)
                    except json.JSONDecodeError:
                        logging.warning(
                            f"Sync Search: Failed JSON decode for {field_name}"
                        )
                        return default_value
                return raw_value if raw_value is not None else default_value

            section_info = safe_json_loads("section_info", [])
            table_info = safe_json_loads("table_info", [])
            figure_info = safe_json_loads("figure_info", [])

            table_references = (
                [str(t) for t in table_refs_raw if t]
                if isinstance(table_refs_raw, list)
                else []
            )
            figure_references = (
                [str(f) for f in figure_refs_raw if f]
                if isinstance(figure_refs_raw, list)
                else []
            )
            section_references = (
                [str(s) for s in section_refs_raw if s]
                if isinstance(section_refs_raw, list)
                else []
            )

            display_name = (
                document_title
                or title
                or (
                    os.path.basename(filepath)
                    if filepath
                    else f"Result {len(processed_references) + 1}"
                )
            )
            reference_parts = []
            if page_number:
                reference_parts.append(f"Page {page_number}")
            if section_references:
                reference_parts.append(f"Sec: {', '.join(section_references)}")
            if table_references:
                reference_parts.append(f"Table: {', '.join(table_references)}")
            if figure_references:
                reference_parts.append(f"Fig: {', '.join(figure_references)}")
            detailed_reference_string = (
                f"{display_name} ({', '.join(reference_parts)})"
                if reference_parts
                else display_name
            )

            ref_key = url or detailed_reference_string

            if ref_key not in processed_references:
                processed_references[ref_key] = {
                    "filename_or_title": detailed_reference_string,
                    "url": url,
                    # Removed "filepath" from response to hide internal blob storage paths
                    "source_type": "document_chunk",
                    "page_number": page_number,
                    "document_title": document_title,
                    "section_references": section_references,
                    "table_references": table_references,
                    "figure_references": figure_references,
                    "section_info": section_info,
                    "table_info": table_info,
                    "figure_info": figure_info,
                    "content": doc_content,
                    "score": doc.get("@search.score"),
                    "reranker_score": doc.get("@search.reranker_score"),
                }
        references_data = list(processed_references.values())

    except Exception as e:
        logging.error(
            f"Error during Sync Azure Search for index '{index_name}': {e}",
            exc_info=True,
        )
        return f"Error accessing sync search index '{index_name}': {e}", []
    finally:
        # Close the potentially created client instance if needed (Sync client might not require explicit close here)
        if "current_search_client" in locals() and hasattr(
            current_search_client, "close"
        ):
            try:
                current_search_client.close()
            except:
                pass  # Ignore close errors

    return context.strip(), references_data


async def get_step_intent_async(step_text: str, client_for_planning: AzureOpenAI):
    """
    Analyzes an analysis step (instruction) and decides which data sources are needed. Async version.
    Returns a dictionary: {"use_live_document": bool, "use_historical_index": bool, "search_query": str|None}
    """
    planning_deployment_id = config.planning_llm_deployment_id  # Get from config
    if not client_for_planning or check_creds(planning_deployment_id):
        logging.warning(
            "Routing LLM client not configured. Defaulting to use both sources."
        )
        return {
            "use_live_document": True,
            "use_historical_index": True,
            "search_query": step_text,
        }

    #     system_prompt = f"""You are a strict query routing assistant for analyzing market analysis documents related to revenue, finance, and market trends.
    # Your job is to decide which data sources are needed for an 'Instruction'.
    # Sources: 'LIVE_DOCUMENT' (the uploaded document), 'HISTORICAL_INDEX' (market reports, financial records, previous analysis, etc.).

    # **CRITICAL RULE 1: CHECK FOR OFF-TOPIC**
    # First, analyze the 'Instruction'. If it is clearly off-topic (e.g., politics, sports, general knowledge, celebrities) and **NOT** related to revenue, finance, market analysis, or trends, you **MUST** respond with:
    # {{"use_live_document": false, "use_historical_index": false, "search_query": "off-topic"}}

    # **CRITICAL RULE 2: IF ON-TOPIC, ROUTE:**
    # - Instruction ONLY about the uploaded document, current revenue, finance analysis, market projections, or specific document content (e.g., "Summarize this doc"):
    #   {{"use_live_document": true, "use_historical_index": false, "search_query": null}}
    # - Instruction involves market comparison, revenue trends, financial history, past performance, or historical financial data:
    #   {{"use_live_document": true, "use_historical_index": true, "search_query": "[Concise query for historical/comparison part]"}}
    # - Instruction ONLY asks for historical/comparison (e.g., "Revenue trends for 2020-2023"):
    #   {{"use_live_document": false, "use_historical_index": true, "search_query": "[Concise query for historical part]"}}

    # If HISTORICAL_INDEX is needed, 'search_query' must be a concise search query.

    # Instruction: "{step_text}"

    # Respond ONLY with a single, valid JSON object."""

    system_prompt = f"""You are a strict but flexible query routing assistant for analyzing instructions related to revenue, finance, company performance, business strategy, and market/industry trends.
Your job is to decide which data sources are needed for an 'Instruction'.
Sources: 'LIVE_DOCUMENT' (the uploaded document), 'HISTORICAL_INDEX' (market reports, industry research, financial history, previous analysis, equity research insights, etc.).

**CRITICAL RULE 1: CHECK FOR OFF-TOPIC**
First, analyze the 'Instruction'. If it is clearly off-topic (e.g., politics, sports, general knowledge, celebrities) and **NOT** related to revenue, finance, business metrics, companies, market trends, or industry analysis, you **MUST** respond with:
{{"use_live_document": false, "use_historical_index": false, "search_query": "off-topic"}}

**CRITICAL RULE 2: IF ON-TOPIC, ROUTE:**
- Instruction ONLY about the uploaded document or content directly inside it (e.g., "Summarize this doc", "Extract key revenue numbers"):
  {{"use_live_document": true, "use_historical_index": false, "search_query": null}}

- Instruction involves company comparison, market comparison, revenue or cost trends, unit economics, historical financial data, industry-wide metrics, or past performance:
  {{"use_live_document": true, "use_historical_index": true, "search_query": "[Concise query for historical/comparison part]"}}

- Instruction ONLY asks for historical, comparative, or industry-wide analysis (e.g., "Compare EBITDA margin between Ather and Ola", "EV market size 2021–2024", "Industry trends for EV 2W"):
  {{"use_live_document": false, "use_historical_index": true, "search_query": "[Concise query for historical/industry part]"}}

If HISTORICAL_INDEX is needed, 'search_query' must be a concise and focused search query.

Instruction: "{step_text}"

Respond ONLY with a single, valid JSON object."""

    try:
        # Use decorated async call for tracing
        @trace_openai_call(client_for_planning, planning_deployment_id, "step_routing")
        async def make_routing_call():
            # Run synchronous SDK call in thread pool
            return await asyncio.to_thread(
                client_for_planning.chat.completions.create,
                model=planning_deployment_id,
                messages=[{"role": "user", "content": system_prompt}],
                response_format={"type": "json_object"},  # Request JSON
                max_tokens=200,  # Limit tokens for routing response
                temperature=0.0,  # Deterministic routing
            )

        response = await make_routing_call()
        intent_data = json.loads(
            response.choices[0].message.content or "{}"
        )  # Default to empty dict

        # Validate structure and types
        if not all(
            k in intent_data
            for k in ["use_live_document", "use_historical_index", "search_query"]
        ):
            raise ValueError("Invalid JSON keys from router")
        if not isinstance(intent_data["use_live_document"], bool) or not isinstance(
            intent_data["use_historical_index"], bool
        ):
            raise ValueError("Invalid boolean types from router")
        if not (
            isinstance(intent_data["search_query"], str)
            or intent_data["search_query"] is None
        ):
            raise ValueError("Invalid search_query type from router")

        # Ensure search_query logic is consistent
        if not intent_data["use_historical_index"]:
            intent_data["search_query"] = None  # Force null if index not used
        elif (
            intent_data["use_historical_index"]
            and not intent_data.get("search_query", "").strip()
        ):
            logging.warning(
                f"Router requested index but provided no query for '{step_text[:50]}...'. Using original step as query."
            )
            intent_data["search_query"] = (
                step_text  # Use original step as fallback query
            )

        logging.info(f"Routing intent for step '{step_text[:50]}...': {intent_data}")
        return intent_data

    except Exception as e:
        logging.error(f"Error in get_step_intent_async: {e}", exc_info=True)
        # Fallback on any error: assume both sources are needed
        return {
            "use_live_document": True,
            "use_historical_index": True,
            "search_query": step_text,
        }


# --- ASYNC CORE RAG Synthesis (NEW, adapted for hybrid context & token protection) ---
async def generate_synthesis_answer(
    user_question: str,
    document_context: str,
    historical_context: str,
    sources: List[Dict],
    max_tokens_param: int,
    client_for_synthesis: AzureOpenAI,
    synthesis_deployment_id: str,
    word_count_target: Optional[int] = None,
    system_prompt_override: Optional[str] = None,
):
    """Generates the final answer using provided context and sources strictly. Async.
    Handles different max token parameter names based on model.
    Conditionally excludes temperature for o3-mini.
    """
    if not client_for_synthesis:
        logging.error("Synthesis LLM client not available.")
        return "Error: Synthesis LLM not configured.", []

    # Determine token limits based on the selected synthesis model
    MODEL_TOKEN_LIMITS = {
        "gpt-4o-mini": 125000,
        "o3-mini": 190000,
    }  # Conservative limits
    max_response_tokens = max(max_tokens_param, 1024)
    MAX_PROMPT_TOKENS = (
        MODEL_TOKEN_LIMITS.get(synthesis_deployment_id, 120000)
        - max_response_tokens
        - 1000
    )

    logging.debug(
        f"Synthesis using {synthesis_deployment_id}. Max prompt tokens: {MAX_PROMPT_TOKENS}, Max response: {max_response_tokens}"
    )

    # --- Construct Final Context String ---
    full_context = ""
    live_doc_part = (
        f"UPLOADED DOCUMENT CONTEXT:\n---------------------\n{document_context}\n---------------------\n\n"
        if document_context
        else ""
    )
    hist_doc_part = (
        f"RETRIEVED HISTORICAL/PEER CONTEXT:\n---------------------\n{historical_context}\n---------------------\n\n"
        if historical_context
        else ""
    )
    full_context = live_doc_part + hist_doc_part

    # --- Prepare Citations (only for historical sources) ---
    unique_sources_list = list(
        {
            (item.get("url") or item.get("filename_or_title")): item
            for item in sources
            if item
        }.values()
    )
    formatted_refs_str = (
        "\n".join(
            [
                f"[doc{i+1}] {html.escape(item['filename_or_title'])}"
                for i, item in enumerate(unique_sources_list)
            ]
        )
        if unique_sources_list
        else "N/A"
    )

    # --- Define Instructions ---
    base_system_prompt = system_prompt_override or SHARED_SYSTEM_PROMPT
    synthesis_instructions = f"""
{base_system_prompt}
**Current Task:** Answer the following USER QUESTION based *only* on the AVAILABLE CONTEXT provided below strictly .
USER QUESTION: {user_question}
AVAILABLE CONTEXT:
=====================
{live_doc_part if live_doc_part else "No uploaded document context provided for this step."}
{hist_doc_part if hist_doc_part else "No historical/peer context provided for this step."}
=====================
IDENTIFIED HISTORICAL/PEER SOURCES (for citation use only):
---------------------
{formatted_refs_str}
---------------------
{f'APPROXIMATE WORD COUNT TARGET: {word_count_target} words.' if word_count_target else ''}
YOUR COMPREHENSIVE, CITED ANSWER (using Markdown format):
"""
    # --- Context Window Protection ---
    synthesis_prompt_for_api = synthesis_instructions
    context_tokens = count_tokens(full_context, synthesis_deployment_id)
    instruction_tokens = count_tokens(
        synthesis_instructions.replace(live_doc_part, "").replace(hist_doc_part, ""),
        synthesis_deployment_id,
    )
    available_context_tokens = MAX_PROMPT_TOKENS - instruction_tokens

    if context_tokens > available_context_tokens:
        logging.warning(
            f"Combined context ({context_tokens} tokens) exceeds limit ({available_context_tokens}) for {synthesis_deployment_id}. Truncating."
        )
        live_tokens = count_tokens(document_context, synthesis_deployment_id)
        hist_tokens = count_tokens(historical_context, synthesis_deployment_id)
        final_live_context, final_hist_context = document_context, historical_context
        if live_tokens + hist_tokens > 0:
            overflow = context_tokens - available_context_tokens
            live_ratio = live_tokens / (live_tokens + hist_tokens)
            hist_ratio = hist_tokens / (live_tokens + hist_tokens)
            live_truncate_by, hist_truncate_by = int(overflow * live_ratio), int(
                overflow * hist_ratio
            )
            if live_truncate_by > 0:
                final_live_context = truncate_text_by_tokens(
                    document_context,
                    live_tokens - live_truncate_by,
                    synthesis_deployment_id,
                )
            if hist_truncate_by > 0:
                final_hist_context = truncate_text_by_tokens(
                    historical_context,
                    hist_tokens - hist_truncate_by,
                    synthesis_deployment_id,
                )
            final_live_doc_part = (
                f"UPLOADED DOCUMENT CONTEXT:\n---------------------\n{final_live_context}\n---------------------\n\n"
                if final_live_context
                else ""
            )
            final_hist_doc_part = (
                f"RETRIEVED HISTORICAL/PEER CONTEXT:\n---------------------\n{final_hist_context}\n---------------------\n\n"
                if final_hist_context
                else ""
            )
            synthesis_prompt_for_api = f"""
{base_system_prompt}
USER QUESTION: {user_question}
AVAILABLE CONTEXT:
=====================
{final_live_doc_part if final_live_doc_part else "No uploaded document context provided (or truncated)."}
{final_hist_doc_part if final_hist_doc_part else "No historical/peer context provided (or truncated)."}
=====================
IDENTIFIED HISTORICAL/PEER SOURCES:\n{formatted_refs_str}\n---------------------\n
{f'APPROX WORD COUNT: {word_count_target}' if word_count_target else ''}\nYOUR ANSWER:\n"""
        else:
            logging.error(
                "Context overflow detected but individual context lengths were zero."
            )
            synthesis_prompt_for_api = synthesis_instructions

    final_prompt_tokens = count_tokens(
        synthesis_prompt_for_api, synthesis_deployment_id
    )
    absolute_limit = MODEL_TOKEN_LIMITS.get(synthesis_deployment_id, 128000)
    if final_prompt_tokens >= absolute_limit:
        logging.error(
            f"FATAL: Final prompt tokens ({final_prompt_tokens}) exceed absolute model limit ({absolute_limit})."
        )
        return (
            f"Error: Input context too large ({final_prompt_tokens} tokens) even after truncation.",
            unique_sources_list,
        )

    # --- Final API Call ---
    try:

        @trace_openai_call(
            client_for_synthesis, synthesis_deployment_id, "synthesis_generation"
        )
        async def make_synthesis_call():
            # --- FIX: Conditionally set parameters ---
            api_params = {
                "model": synthesis_deployment_id,
                "messages": [{"role": "user", "content": synthesis_prompt_for_api}],
            }
            # o3-mini expects 'max_completion_tokens' and does NOT support 'temperature'
            if "o3-mini" in synthesis_deployment_id:
                api_params["max_completion_tokens"] = max_response_tokens
                # Do NOT add temperature
            # Other models (like gpt-4o-mini) expect 'max_tokens' and support 'temperature'
            else:
                api_params["max_tokens"] = max_response_tokens
                api_params["temperature"] = 0.2  # Or get from config/state
            # --- End FIX ---

            # Run synchronous SDK call in thread pool
            return await asyncio.to_thread(
                client_for_synthesis.chat.completions.create,
                **api_params,  # Pass the constructed parameters
            )

        response = await make_synthesis_call()
        final_answer = response.choices[0].message.content or "[No content generated]"
        logging.debug(f"Synthesis successful. Output length: {len(final_answer)}")
        return final_answer, unique_sources_list
    except Exception as e:
        logging.error(f"Error during synthesis API call: {e}", exc_info=True)
        # Check if the error is the specific BadRequestError for temperature
        if isinstance(e, openai.BadRequestError) and "temperature" in str(e):
            logging.error(
                f"Temperature parameter likely unsupported for model {synthesis_deployment_id}."
            )
        return f"Error generating synthesized answer: {e}", unique_sources_list


# --- End Async Helpers ---


# --- TOKEN COUNTING UTILITIES (NEW) ---
@functools.lru_cache(maxsize=128)
def get_encoding_for_model(model_name: str):
    """Gets the tiktoken encoding for a model."""
    try:
        return tiktoken.encoding_for_model(model_name)
    except KeyError:
        return tiktoken.get_encoding("cl100k_base")


def count_tokens(text: str, model_name: str) -> int:
    """Counts tokens."""
    if not text:
        return 0
    # Handle potential non-string input gracefully
    if not isinstance(text, str):
        text = str(text)
    return len(get_encoding_for_model(model_name).encode(text))


def truncate_text_by_tokens(text: str, max_tokens: int, model_name: str) -> str:
    """Truncates text to token limit (simple head truncation)."""
    if not text:
        return ""
    # Handle potential non-string input gracefully
    if not isinstance(text, str):
        text = str(text)
    encoding = get_encoding_for_model(model_name)
    tokens = encoding.encode(text)
    if len(tokens) > max_tokens:
        logging.warning(f"Truncating text from {len(tokens)} to {max_tokens} tokens.")
        # Simple head truncation - consider middle truncation if needed
        # Use decode with error handling if necessary, though simple truncation is usually safe
        return encoding.decode(tokens[:max_tokens])
    return text

async def query_gemini_deep_research(question: str, api_key: str) -> str:
    """
    Runs Gemini Deep Research and waits until completion.
    WARNING: This can take 10–45 minutes.
    """
    client = genai.Client(api_key=api_key)

    interaction = client.interactions.create(
        input=question,
        agent="deep-research-pro-preview-12-2025",
        background=True,  # Required for long-running jobs
    )

    start_time = time.time()

    while True:
        interaction = client.interactions.get(interaction.id)

        if interaction.status == "completed":
            return interaction.outputs[-1].text

        if interaction.status == "failed":
            return f"Error: {getattr(interaction, 'error', 'Unknown error')}"

        # Non-blocking sleep for async apps
        await asyncio.sleep(20)
        
        
async def query_serpapi(query: str, count: int = 10) -> str:
    """MODIFICATION: Performs a web search using httpx for non-blocking requests."""
    try:
        if not query or not isinstance(query, str) or not query.strip():
            logging.error(f"[SerpApi] Invalid or empty query: {query!r}")
            return "Error: Invalid or empty query provided."

        if check_creds(config.serpapi_api_key):
            logging.warning("[SerpApi] API key not configured.")
            return "Error: SerpApi API key not configured."

        params = {"q": query.strip(), "api_key": config.serpapi_api_key, "num": count}
        url = "https://serpapi.com/search.json"

        logging.info(f"[SerpApi] Sending request: {url}")
        response = await clients.async_http_client.get(url, params=params)
        response.raise_for_status()
        data = response.json()

        organic_results = data.get("organic_results", [])
        if not organic_results:
            return "Error: No web search results found via SerpApi."

        results_text = "\n".join(
            [
                f"Title: {res.get('title', 'N/A')}\nSnippet: {res.get('snippet', 'N/A')}"
                for res in organic_results
            ]
        )
        return results_text or "Error: No usable search results returned."
    except httpx.HTTPStatusError as e:
        logging.error(
            f"[SerpApi] HTTP error: {e.response.status_code} for query: {query}"
        )
        return f"Error during SerpApi request: HTTP {e.response.status_code}"
    except Exception as e:
        logging.error(f"[SerpApi] Unexpected error: {e}", exc_info=True)
        return f"Error: Unexpected issue in SerpApi query ({type(e).__name__})"


async def query_bing_web_search(query: str, count: int = 5) -> str:
    """MODIFICATION: Uses httpx for non-blocking Bing search requests."""
    if check_creds(config.bing_search_api_key):
        return "Error: Bing Search API key not configured."

    headers = {"Ocp-Apim-Subscription-Key": config.bing_search_api_key}
    params = {"q": query, "count": count}

    try:
        response = await clients.async_http_client.get(
            config.bing_search_endpoint, headers=headers, params=params
        )
        response.raise_for_status()
        data = response.json()
        snippets = [
            f"Title: {res['name']}\nSnippet: {res['snippet']}"
            for res in data.get("webPages", {}).get("value", [])
        ]
        return "\n".join(snippets) if snippets else "No web search results found."
    except httpx.HTTPStatusError as e:
        logging.error(
            f"Error during Bing web search: {e.response.status_code}", exc_info=True
        )
        return f"Error during Bing web search: HTTP {e.response.status_code}"
    except Exception as e:
        logging.error(f"Error during Bing web search: {e}", exc_info=True)
        return f"Error during Bing web search: {type(e).__name__}"


# --- ASYNC Query Planning (Corrected for Chat RAG) ---
async def get_query_plan_from_llm(
    user_question: str, history: List[ChatMessage]
) -> tuple[Optional[str], List[str]]:
    """
    Breaks down a user question into search queries using an LLM. Async version.
    Always returns a tuple: (error_message_or_None, list_of_queries).
    """
    if not clients.planning_openai_client:
        logging.warning(
            "Planning LLM client not configured. Returning single-step plan."
        )
        # Return error message and fallback query list
        return "Warning: Planning LLM client not available", [user_question]

    history_str = "\n".join([f"{msg.role}: {msg.content}" for msg in history])
    # Ensure consistent prompt structure with clear JSON instruction
    planning_prompt = f"""You are a query planning assistant for analyzing AERA regulatory documents (DIAL, MIAL, BIAL, HIAL).
Break down the user's complex question into 1-10 simple, self-contained search queries for an index.
Return ONLY a valid JSON list of strings (the queries). No explanations.

CONVERSATION HISTORY (if any):
{history_str}

USER QUESTION: {user_question}

JSON list of search queries:"""

    try:
        # Use decorated async call for tracing
        @trace_openai_call(
            clients.planning_openai_client,
            config.planning_llm_deployment_id,
            "query_planning",
        )
        async def make_planning_call():
            # Run synchronous SDK call in thread pool
            return await asyncio.to_thread(
                clients.planning_openai_client.chat.completions.create,  # Corrected attribute name
                model=config.planning_llm_deployment_id,
                messages=[{"role": "user", "content": planning_prompt}],
                response_format={
                    "type": "json_object"
                },  # Request JSON directly if model supports it
                temperature=0.0,  # Deterministic planning
                max_tokens=500,  # Reduced max tokens for plan
            )

        response = await make_planning_call()
        plan_str = (
            response.choices[0].message.content or "[]"
        )  # Default to empty list string

        # Attempt to parse the JSON directly
        try:
            query_plan = json.loads(plan_str)
            # Validate structure: list of strings
            if isinstance(query_plan, dict) and "queries" in query_plan:
                query_plan = query_plan["queries"]

            if (
                isinstance(query_plan, list)
                and all(isinstance(q, str) for q in query_plan)
                and query_plan
            ):
                logging.info(f"Successfully planned {len(query_plan)} queries.")
                return None, query_plan  # Success: No error, return plan
            else:
                # Handle cases like empty list or invalid structure within JSON
                logging.warning(
                    f"Planner returned invalid JSON structure or empty list: {plan_str}. Falling back."
                )
                return None, [
                    user_question
                ]  # Success (fallback): No error, return original question

        except json.JSONDecodeError:
            # Fallback if the response wasn't valid JSON (despite asking for it)
            logging.warning(
                f"Planner response was not valid JSON: {plan_str}. Falling back."
            )
            # Still return None for error, just use original question
            return None, [
                user_question
            ]  # Success (fallback): No error, return original question

    except Exception as e:
        # Catch any other exceptions during API call or processing
        logging.error(f"Error during query planning: {e}", exc_info=True)
        # Return the error message and the fallback query list
        return f"Error during planning: {str(e)}", [user_question]


# --- ORIGINAL RAG Pipeline (Corrected for Chat - uses sync search via threads) ---
async def execute_advanced_rag_pipeline(
    user_question: str, history: Optional[List[ChatMessage]] = None
):
    """Original RAG pipeline using planning and SYNC index search only (for chat)."""
    if not clients.synthesis_openai_client or not clients.planning_openai_client:
        raise HTTPException(status_code=503, detail="AI services not configured.")
    history = history or []

    logging.info(f"Chat RAG (Sync Search): Planning for '{user_question[:100]}...'")
    # Make sure get_query_plan_from_llm is correctly defined and awaited
    plan_error, query_plan = await get_query_plan_from_llm(user_question, history)
    if plan_error:
        logging.warning(f"Chat RAG: Planning error: {plan_error}. Using single step.")
        # Fallback to single query if planning failed
        query_plan = [user_question]
    elif not query_plan:  # Ensure query_plan is never empty
        logging.warning(f"Chat RAG: Planner returned empty plan. Using single step.")
        query_plan = [user_question]

    logging.info(
        f"Chat RAG (Sync Search): Executing {len(query_plan)} search steps sequentially..."
    )

    combined_context = ""
    all_retrieved_details = []  # Use this to collect sources from all steps

    # --- CORRECTED SEARCH EXECUTION ---
    # Process each sub_query one by one, awaiting the result before the next
    for i, sub_query in enumerate(query_plan):
        try:
            # Run sync search in thread for the current sub_query
            logging.debug(
                f"Chat RAG: Step {i+1} - Calling query_azure_search_sync for '{sub_query[:50]}...'"
            )
            # *** THIS IS THE CORRECTED CALL ***
            context_for_step, retrieved_details_for_step = await asyncio.to_thread(
                query_azure_search_sync,  # *** Call the SYNC version ***
                sub_query,
                config.default_search_index_name,
                k=5,  # Number of results per step
                use_hybrid_semantic_search=True,  # Assuming hybrid for chat RAG
            )
            logging.debug(f"Chat RAG: Step {i+1} - query_azure_search_sync returned.")

            # Process results for this step
            if context_for_step and not context_for_step.startswith("Error"):
                logging.info(
                    f"Chat RAG: Step {i+1} - Retrieved {len(retrieved_details_for_step)} sources."
                )
                combined_context += (
                    f"\n\n--- Context for sub-query: '{sub_query}' ---\n"
                    + context_for_step
                )
                all_retrieved_details.extend(retrieved_details_for_step)
            else:
                # Log if no results or if search returned an error string
                logging.warning(
                    f"Chat RAG (Sync Search): No results or error for step '{sub_query}': {context_for_step}"
                )
        except Exception as search_err:
            # Log error for the specific step but continue processing other steps
            logging.error(
                f"Chat RAG (Sync Search): Unhandled exception during search step '{sub_query}': {search_err}",
                exc_info=True,
            )
            # Optionally add a note to the context or skip? For now, just log and continue.
    # --- END OF CORRECTION ---

    # --- Synthesis part ---
    if not combined_context.strip():
        logging.info(
            "Chat RAG (Sync Search): No relevant context found after all search steps."
        )
        # Return empty sources list here
        return {
            "answer": "I couldn't find relevant information in the knowledge base to answer that.",
            "plan": query_plan,
            "sources": [],
        }

    # Use the async synthesis function (passing only historical context)
    logging.info("Chat RAG (Sync Search): Synthesizing final answer...")
    # Ensure generate_synthesis_answer is correctly defined and awaited
    final_answer, unique_sources = await generate_synthesis_answer(
        user_question=user_question,
        document_context="",  # No live document for chat RAG
        historical_context=combined_context.strip(),
        sources=all_retrieved_details,  # Pass all aggregated sources for citation generation
        max_tokens_param=config.max_output_tokens,  # Use configured max output
        client_for_synthesis=clients.synthesis_openai_client,
        synthesis_deployment_id=config.deployment_id,  # Use default chat model (deployment_id)
    )

    logging.info(
        f"Chat RAG (Sync Search): Synthesis complete. Answer length: {len(final_answer)}, Sources cited: {len(unique_sources)}"
    )
    # Return unique sources derived by generate_synthesis_answer
    return {"answer": final_answer, "plan": query_plan, "sources": unique_sources}


def parse_html_to_docx(soup, document):
    """A robust function to parse HTML content and convert it into a .docx document."""
    content_root = soup.body if soup.body else soup
    for element in content_root.children:
        if isinstance(element, NavigableString):
            if element.string and element.string.strip():
                document.add_paragraph(element.string.strip())
            continue
        if hasattr(element, "name") and element.name:
            if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                try:
                    level = int(element.name[1])
                    document.add_heading(element.get_text(strip=True), level=level)
                except (ValueError, IndexError):
                    document.add_heading(element.get_text(strip=True), level=2)
            elif element.name == "p":
                p = document.add_paragraph()
                for content in element.contents:
                    if hasattr(content, "name") and content.name in ["b", "strong"]:
                        p.add_run(content.get_text(strip=True)).bold = True
                    elif hasattr(content, "name") and content.name in ["i", "em"]:
                        p.add_run(content.get_text(strip=True)).italic = True
                    else:
                        p.add_run(str(content))
            elif element.name in ["ul", "ol"]:
                style = "List Bullet" if element.name == "ul" else "List Number"
                for li in element.find_all("li", recursive=False):
                    document.add_paragraph(li.get_text(strip=True), style=style)
            elif element.name == "table":
                rows = element.find_all("tr")
                if not rows:
                    continue
                header_rows, body_rows = element.select("thead > tr"), element.select(
                    "tbody > tr"
                )
                if not header_rows and not body_rows:
                    header_rows, body_rows = rows[0:1], rows[1:]
                elif not body_rows:
                    body_rows = []
                header_cells = (
                    header_rows[0].find_all(["th", "td"]) if header_rows else []
                )
                if not header_cells:
                    continue
                table = document.add_table(rows=0, cols=len(header_cells))
                table.style = "Table Grid"
                table.autofit = True
                for h_row in header_rows:
                    cells, row_cells = (
                        h_row.find_all(["th", "td"]),
                        table.add_row().cells,
                    )
                    for i, cell in enumerate(cells):
                        if i < len(row_cells):
                            p = row_cells[i].paragraphs[0]
                            p.add_run(cell.get_text(strip=True)).bold = True
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for b_row in body_rows:
                    cells = b_row.find_all("td")
                    if len(cells) == len(header_cells):
                        row_cells = table.add_row().cells
                        for i, cell in enumerate(cells):
                            row_cells[i].text = cell.get_text(strip=True)
            elif element.name == "hr":
                document.add_page_break()


# --- API ENDPOINTS ---
@app.get("/")
async def read_root():  # Use async def
    logging.info("Root endpoint accessed")
    return {"message": "BIAL Regulatory Platform MDA API is running."}


@app.get("/health/app-insights")
def app_insights_health_check():
    """Health check endpoint to verify Azure Monitor OpenTelemetry configuration."""
    logging.info("Azure Monitor health check endpoint accessed")
    config_status = {
        "APPLICATIONINSIGHTS_CONNECTION_STRING": bool(
            APPLICATIONINSIGHTS_CONNECTION_STRING
        ),
        "LOG_ANALYTICS_WORKSPACE_ID": bool(LOG_ANALYTICS_WORKSPACE_ID),
        "LOG_ANALYTICS_WORKSPACE_NAME": bool(LOG_ANALYTICS_WORKSPACE_NAME),
        "APPLICATION_INSIGHTS_NAME": APPLICATION_INSIGHTS_NAME,
        "azure_monitor_enabled": azure_monitor_enabled,
    }
    is_configured = (
        bool(APPLICATIONINSIGHTS_CONNECTION_STRING) and azure_monitor_enabled
    )
    logging.info(f"Azure Monitor configuration status: {config_status}")
    return {
        "status": "healthy" if is_configured else "misconfigured",
        "azure_monitor_configured": is_configured,
        "configuration": config_status,
        "message": (
            "Azure Monitor OpenTelemetry is properly configured"
            if is_configured
            else "Azure Monitor OpenTelemetry configuration is missing or incomplete"
        ),
    }


@app.post("/register", response_model=User, summary="Register a new user")
async def register(user: UserCreate):
    if not clients.mongo_client:
        raise HTTPException(503, "DB not configured.")
    if await get_user_from_db(user.username):
        raise HTTPException(400, "Username already registered.")
    hashed_password = get_password_hash(user.password)
    doc = {
        "id": user.id or f"user_{uuid.uuid4().hex[:8]}",
        "username": user.username,
        "password_hash": hashed_password,
        "role": user.role or "user",
        "isActive": True if user.isActive is None else bool(user.isActive),
        "createdAt": user.createdAt or datetime.now(timezone.utc),
    }
    if user.lastLogin:
        doc["lastLogin"] = user.lastLogin
    await asyncio.to_thread(
        clients.mongo_client[config.cosmos_database_name][
            config.cosmos_users_collection
        ].insert_one,
        doc,
    )
    return User(
        username=user.username, role=doc.get("role"), isActive=doc.get("isActive", True)
    )


# In mda.py ONLY

# In mda.py


@app.post("/token", response_model=Token, summary="User login to get a token")
async def login(form: OAuth2PasswordRequestForm = Depends()):
    user = await get_user_from_db(form.username)
    if not user or not verify_password(form.password, user.hashed_password):
        raise HTTPException(
            401, "Incorrect username or password", {"WWW-Authenticate": "Bearer"}
        )

    if hasattr(user, "isActive") and not user.isActive:
        logging.warning(f"Attempted login by inactive user: {form.username}")
        raise HTTPException(status_code=401, detail="User account is disabled.")

    # --- FINAL FIX ---
    # Get the login time ONCE and truncate microseconds
    login_time = datetime.now(timezone.utc).replace(microsecond=0)
    # --- END FIX ---

    try:
        if clients.mongo_client:
            await asyncio.to_thread(
                clients.mongo_client[config.cosmos_database_name][
                    config.cosmos_users_collection
                ].update_one,
                {"username": user.username},
                {
                    "$set": {
                        "lastLogin": login_time,
                        "min_token_iat": login_time,  # <-- Use truncated timestamp
                    }
                },
            )
    except Exception:
        pass

    # Pass the SAME truncated timestamp into the token
    token = create_access_token({"sub": user.username}, iat_time=login_time)

    return {"access_token": token, "token_type": "bearer"}


HARDCODED_USERNAME = "raushanpandey@kpmg.com"
HARDCODED_PASSWORD = "raushan123"


@app.post(
    "/login", response_model=Token, summary="User login with username and password"
)
async def login_with_credentials(request: UserCreate):

    # 🔥 TEMP HARDCODED LOGIN (DEV ONLY)
    if (
        request.username == HARDCODED_USERNAME
        and request.password == HARDCODED_PASSWORD
    ):
        login_time = datetime.now(timezone.utc).replace(microsecond=0)

        token = create_access_token(
            {"sub": request.username},
            iat_time=login_time,
        )

        return {
            "access_token": token,
            "token_type": "bearer",
        }

    # ---- NORMAL DB LOGIN FLOW (unchanged) ----
    user = await get_user_from_db(request.username)
    if not user or not verify_password(request.password, user.hashed_password):
        raise HTTPException(
            401, "Incorrect username or password", {"WWW-Authenticate": "Bearer"}
        )

    if hasattr(user, "isActive") and not user.isActive:
        raise HTTPException(status_code=401, detail="User account is disabled.")

    login_time = datetime.now(timezone.utc).replace(microsecond=0)

    try:
        if clients.mongo_client:
            await asyncio.to_thread(
                clients.mongo_client[config.cosmos_database_name][
                    config.cosmos_users_collection
                ].update_one,
                {"username": user.username},
                {
                    "$set": {
                        "lastLogin": login_time,
                        "min_token_iat": login_time,
                    }
                },
            )
    except Exception:
        pass

    token = create_access_token({"sub": user.username}, iat_time=login_time)

    return {"access_token": token, "token_type": "bearer"}


@app.post("/logout", summary="Logout current user and invalidate token")
async def logout_current_user(token: str = Depends(oauth2_scheme)):
    global clients
    if not clients.mongo_client:
        raise HTTPException(503, "DB not configured for logout.")
    try:
        # Decode the token to get its ID (jti) and expiry (exp)
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        jti = payload.get("jti")
        exp = payload.get("exp")

        if jti and exp:
            # Store the JTI in a new 'token_blocklist' collection
            # IMPORTANT: You must enable TTL on this collection in Cosmos DB
            blocklist_collection = clients.mongo_client[config.cosmos_database_name][
                "token_blocklist"
            ]
            await asyncio.to_thread(
                blocklist_collection.insert_one,
                {
                    "jti": jti,
                    # Store expiry time as a BSON Date for Cosmos DB TTL
                    "expiry_time": datetime.fromtimestamp(exp, timezone.utc),
                },
            )
            logging.info(f"Token {jti} has been blocklisted (revoked).")

        return {"status": "success", "message": "Logged out and token revoked."}
    except JWTError as e:
        raise HTTPException(status_code=401, detail=f"Invalid token for logout: {e}")


def require_admin(request: Request):
    if not ADMIN_API_KEY:
        raise HTTPException(status_code=503, detail="Admin API not configured")
    provided = request.headers.get("x-admin-key") or request.headers.get("X-Admin-Key")
    if provided != ADMIN_API_KEY:
        raise HTTPException(status_code=401, detail="Unauthorized")


@app.post("/admin/create-user", summary="Admin: create user with full schema")
async def admin_create_user(req: AdminCreateUserRequest, request: Request):
    require_admin(request)
    if not clients.mongo_client:
        raise HTTPException(503, "DB not configured.")
    if await get_user_from_db(req.username):
        raise HTTPException(400, "Username already registered.")
    password_hash = req.password_hash or (
        get_password_hash(req.password) if req.password else None
    )
    if not password_hash:
        raise HTTPException(400, "Either password or password_hash is required")
    doc = {
        "id": req.id or f"user_{uuid.uuid4().hex[:8]}",
        "username": req.username,
        "password_hash": password_hash,
        "role": req.role or "user",
        "isActive": req.isActive if req.isActive is not None else True,
        "createdAt": req.createdAt or datetime.now(timezone.utc),
    }
    if req.lastLogin:
        doc["lastLogin"] = req.lastLogin
    await asyncio.to_thread(
        clients.mongo_client[config.cosmos_database_name][
            config.cosmos_users_collection
        ].insert_one,
        doc,
    )
    return {"status": "created", "username": req.username, "id": doc["id"]}


@app.post("/maintenance/set-password", summary="Maintenance: set password for a user")
async def maintenance_set_password(req: MaintenanceSetPasswordRequest):
    if not ALLOW_MAINTENANCE:
        raise HTTPException(403, "Maintenance operations are disabled")
    if not clients.mongo_client:
        raise HTTPException(503, "DB not configured.")
    if not await get_user_from_db(req.username):
        raise HTTPException(404, "User not found")
    hashed = get_password_hash(req.new_password)
    await asyncio.to_thread(
        clients.mongo_client[config.cosmos_database_name][
            config.cosmos_users_collection
        ].update_one,
        {"username": req.username},
        {"$set": {"password_hash": hashed}},
    )
    return {"status": "updated", "username": req.username}


@app.post("/login/sso", response_model=Token, summary="Login via Azure Entra ID SSO")
async def sso_login(req: SSOLoginRequest):
    logging.info("SSO login attempt received")
    if DEV_MODE:
        logging.info("DEV_MODE enabled, bypassing SSO validation")
        return {
            "access_token": create_access_token(data={"sub": "dev_user"}),
            "token_type": "bearer",
        }

    if not all([TENANT_ID, CLIENT_ID, clients.jwks_client]):
        logging.error(
            f"SSO configuration missing: TENANT_ID={bool(TENANT_ID)}, CLIENT_ID={bool(CLIENT_ID)}, jwks_client={bool(clients.jwks_client)}"
        )
        raise HTTPException(503, "SSO service not configured correctly.")

    try:
        signing_key = clients.jwks_client.get_signing_key_from_jwt(req.sso_token).key
        decoded_token = pyjwt.decode(
            req.sso_token,
            signing_key,
            algorithms=["RS256"],
            audience=CLIENT_ID,
            issuer=ISSUER_URL,
        )
        username = (
            decoded_token.get("preferred_username")
            or decoded_token.get("upn")
            or decoded_token.get("email")
            or decoded_token.get("sub")
        )
        username = username.lower()
        if not username:
            raise HTTPException(400, "Username not found in SSO token.")

        logging.info(f"SSO login successful for user: {username}")
        access_token = create_access_token(data={"sub": username})
        return {"access_token": access_token, "token_type": "bearer"}
    except pyjwt.ExpiredSignatureError:
        raise HTTPException(401, "SSO token has expired.")
    except pyjwt.InvalidAudienceError as e:
        raise HTTPException(401, f"Invalid SSO token audience: {e}")
    except pyjwt.InvalidIssuerError as e:
        raise HTTPException(401, f"Invalid SSO token issuer: {e}")
    except Exception as e:
        logging.error(f"Could not validate SSO token: {e}", exc_info=True)
        raise HTTPException(401, f"Could not validate SSO token: {e}")


@app.get("/user/me", response_model=User, summary="Get current user information")
async def get_current_user_info(current_user: User = Depends(auth_dependency)):
    # The 'auth_dependency' (which runs 'get_current_user_prod')
    # has already done all the work.
    # It has either returned a valid, active user or raised a 401 error.
    # We just need to return the user object it provided.
    return current_user


async def stream_analysis_process_enhanced(
    extracted_text: str,
    prompt_template: str,
    model_deployment_id: str,
    word_count: int,
    current_user: User,
    analysis_title: str,
):
    """Async generator for streaming analysis results with routing."""
    start_time = time.time()
    final_report_for_log = "Error during generation"  # Default log message
    tracer = trace.get_tracer(__name__)
    parent_span = trace.get_current_span()  # Get current request span if available

    # Start a new span for the entire streaming process, linked to the request span
    with tracer.start_as_current_span(
        "streaming_analysis_enhanced", context=trace.set_span_in_context(parent_span)
    ) as stream_span:
        try:
            # Set attributes for the overall process
            stream_span.set_attribute("analysis.title", analysis_title)
            stream_span.set_attribute("analysis.model", model_deployment_id)
            # ... (other attributes) ...
            stream_span.set_attribute("user.id", current_user.username)

            # 1. Parse Steps
            analysis_steps = [
                s.strip() for s in re.findall(r"^\s*\d+-\s*(.*)", prompt_template, re.M)
            ] or [prompt_template.strip()]
            stream_span.set_attribute("analysis.num_steps", len(analysis_steps))
            logging.info(
                f"Starting analysis '{analysis_title}' with {len(analysis_steps)} steps for user {current_user.username}"
            )

            step_results = []
            all_sources_combined = (
                []
            )  # Aggregate sources from all steps for final report

            # --- [START] NEW VARIABLE ---
            off_topic_step_count = 0
            # --- [END] NEW VARIABLE ---

            # 2. Process Each Step with Routing
            for i, step_text in enumerate(analysis_steps):
                step_start_time = time.time()
                # Start a child span for each step
                with tracer.start_as_current_span(f"analysis_step_{i+1}") as step_span:
                    # ... (set attributes, logging, yield step_start) ...
                    step_span.set_attribute("step.number", i + 1)
                    step_span.set_attribute("step.instruction", step_text[:500])
                    logging.info(
                        f"Step {i+1}/{len(analysis_steps)}: Routing '{step_text[:100]}...'"
                    )
                    yield f"data: {json.dumps({'type': 'step_start', 'step': i+1, 'total': len(analysis_steps), 'title': step_text})}\n\n"

                    # Get Intent (await the async router)
                    intent = await get_step_intent_async(
                        step_text, clients.planning_openai_client
                    )
                    logging.info(f"DEBUG Step {i+1} Intent: {intent}")  # DEBUG LOG

                    # --- GUARDRAIL BLOCK (MODIFIED) ---
                    # Check if the router flagged this step as off-topic
                    if intent.get("search_query") == "off-topic" or (
                        not intent.get("use_live_document")
                        and not intent.get("use_historical_index")
                        and intent.get("search_query") is None
                    ):
                        logging.warning(
                            f"Step {i+1} blocked as off-topic by router for prompt: '{step_text[:100]}...'"
                        )
                        step_span.set_attribute("step.intent.off_topic", True)

                        # --- [START] INCREMENT COUNTER ---
                        off_topic_step_count += 1
                        # --- [END] INCREMENT COUNTER ---

                        # Define the refusal message
                        step_answer = "This question is off-topic. I can only provide analysis related to the uploaded document or regulatory matters."

                        # Yield the refusal to the user
                        yield f"data: {json.dumps({'type': 'step_result', 'step': i+1, 'answer': step_answer, 'sources': []})}\n\n"

                        # Store this refusal in the final report results
                        step_results.append(
                            {
                                "step": step_text,
                                "answer": step_answer,
                                "sources": [],
                            }
                        )
                        # Skip to the next step in the loop
                        continue
                    # --- [END] GUARDRAIL BLOCK ---

                    # ... (rest of the step processing: set attributes, prepare context, search block) ...
                    # ... (This part is unchanged) ...
                    step_span.set_attribute(
                        "step.intent.use_live_document", intent.get("use_live_document")
                    )
                    # ... (etc.) ...

                    # --- [START] MODIFIED SEARCH BLOCK ---
                    step_live_context = (
                        extracted_text if intent.get("use_live_document") else ""
                    )
                    step_historical_context = ""
                    step_sources = []

                    if intent.get("use_historical_index") and intent.get(
                        "search_query"
                    ):
                        search_query = intent["search_query"]
                        logging.info(
                            f"DEBUG Step {i+1}: Attempting index search with query: '{search_query}'"
                        )
                        yield f"data: {json.dumps({'type': 'status_update', 'step': i+1, 'message': f'Searching index: {search_query[:100]}...'})}\n\n"

                        step_historical_context, step_sources = (
                            await query_azure_search_async(
                                search_query,
                                config.default_search_index_name,
                                k=5,
                                use_hybrid_semantic=True,
                            )
                        )
                        # ... (rest of search logging/handling) ...
                        if step_historical_context.startswith("Error"):
                            logging.error(
                                f"DEBUG Step {i+1}: Index search returned error: {step_historical_context}"
                            )
                            yield f"data: {json.dumps({'type': 'step_error', 'step': i+1, 'message': f'Index Search Failed: {step_historical_context}'})}\n\n"
                            step_historical_context = ""
                            step_sources = []
                        elif not step_sources:
                            logging.warning(
                                f"DEBUG Step {i+1}: Index search returned NO sources for query '{search_query}'."
                            )
                        else:
                            logging.info(
                                f"DEBUG Step {i+1}: Index search successful, retrieved {len(step_sources)} sources."
                            )
                            all_sources_combined.extend(step_sources)
                            step_span.set_attribute(
                                "step.retrieved_historical_chunks", len(step_sources)
                            )

                    elif intent.get("use_historical_index"):
                        logging.warning(
                            f"DEBUG Step {i+1}: Router needed index, but no query generated."
                        )
                        yield f"data: {json.dumps({'type': 'warning', 'step': i+1, 'message': 'Routing suggested index search, but no query was formed.'})}\n\n"
                    else:
                        logging.info(
                            f"DEBUG Step {i+1}: Skipping index search based on routing intent."
                        )
                    # --- [END] MODIFIED SEARCH BLOCK ---

                    # ... (Generate Answer for Step - This part is unchanged) ...
                    yield f"data: {json.dumps({'type': 'status_update', 'step': i+1, 'message': 'Generating analysis...'})}\n\n"
                    logging.info(f"Step {i+1}: Synthesizing answer...")
                    step_answer, sources_for_step_display = (
                        await generate_synthesis_answer(
                            user_question=step_text,
                            document_context=step_live_context,
                            historical_context=step_historical_context,
                            sources=step_sources,
                            max_tokens_param=config.max_output_tokens,
                            client_for_synthesis=clients.synthesis_openai_client,
                            synthesis_deployment_id=model_deployment_id,
                        )
                    )
                    # ... (rest of step result handling - unchanged) ...
                    step_span.set_attribute("step.answer_length", len(step_answer))
                    step_duration_ms = round((time.time() - step_start_time) * 1000)
                    step_span.set_attribute("duration.ms", step_duration_ms)
                    logging.info(
                        f"Step {i+1}: Synthesis complete in {step_duration_ms} ms."
                    )

                    if step_answer.startswith("Error"):
                        step_span.set_status(Status(StatusCode.ERROR, step_answer))
                        yield f"data: {json.dumps({'type': 'step_error', 'step': i+1, 'message': step_answer})}\n\n"
                        step_results.append(
                            {
                                "step": step_text,
                                "answer": f"*Error processing this step: {step_answer}*",
                                "sources": [],
                            }
                        )
                    else:
                        step_span.set_status(Status(StatusCode.OK))
                        yield f"data: {json.dumps({'type': 'step_result', 'step': i+1, 'answer': step_answer, 'sources': sources_for_step_display})}\n\n"
                        step_results.append(
                            {
                                "step": step_text,
                                "answer": step_answer,
                                "sources": sources_for_step_display,
                            }
                        )

                    await asyncio.sleep(0.05)

            # --- [START] NEW REFINEMENT GUARDRAIL ---

            # 3. Final Refinement (Now with a guardrail)

            # Check if all steps were flagged as off-topic
            if off_topic_step_count == len(analysis_steps):
                logging.info("All steps were off-topic. Skipping final refinement.")

                # Use the refusal message from the first (and only) step as the final report
                final_report = step_results[0]["answer"]
                final_report_for_log = final_report

                # Ensure the final sources list is empty
                unique_final_sources = []

            else:
                # --- This is the ORIGINAL refinement code, now in the ELSE block ---
                logging.info("Starting final report refinement...")
                yield f"data: {json.dumps({'type': 'refine_start', 'message': 'Combining and refining final report...'})}\n\n"
                refine_start_time = time.time()

                with tracer.start_as_current_span("final_refinement") as refine_span:
                    # Prepare combined text from step results
                    combined_step_text = "\n\n---\n\n".join(
                        [
                            f"### Step {i+1}: {res['step']}\n{res['answer']}"
                            for i, res in enumerate(step_results)
                        ]
                    )

                    refinement_prompt = f"""Synthesize the following analyzed sections into a single, cohesive report (approx {word_count} words). Ensure logical flow, remove redundancies, format professionally using Markdown. Address any noted errors if possible, otherwise summarize briefly or omit the problematic step's content. Also, make sure to combine all conclusion/summary sections into a single conclusion/summary section, and strictly remove all other conclusion/summary sections.
                    ANALYZED SECTIONS:\n---\n{combined_step_text}\n---\nFINAL COHESIVE REPORT (Markdown Format):"""

                    # Call synthesis for refinement (await async)
                    final_report, _ = await generate_synthesis_answer(
                        user_question="Synthesize the final report from analyzed sections.",
                        document_context="",
                        historical_context="",
                        sources=[],  # No direct context/sources for refinement
                        max_tokens_param=max(
                            int(word_count * 2.0), config.max_output_tokens
                        ),  # Generous token limit
                        client_for_synthesis=clients.synthesis_openai_client,
                        synthesis_deployment_id=model_deployment_id,  # Use same model
                        word_count_target=word_count,
                        system_prompt_override=refinement_prompt,  # Use the specific refinement prompt
                    )
                    final_report_for_log = final_report  # Capture for logging

                    refine_duration_ms = round((time.time() - refine_start_time) * 1000)
                    refine_span.set_attribute("duration.ms", refine_duration_ms)
                    refine_span.set_attribute(
                        "refinement.final_length", len(final_report)
                    )
                    logging.info(
                        f"Final refinement complete in {refine_duration_ms} ms."
                    )

                    # Use aggregated sources from ALL steps for the final report's citation list
                    unique_final_sources = list(
                        {
                            (s.get("url") or s.get("filename_or_title")): s
                            for s in all_sources_combined
                            if s
                        }.values()
                    )
            # --- [END] NEW REFINEMENT GUARDRAIL ---

            stream_span.set_attribute("analysis.final_report_length", len(final_report))
            stream_span.set_attribute(
                "analysis.total_sources_aggregated", len(unique_final_sources)
            )
            stream_span.set_status(
                Status(StatusCode.OK)
            )  # Mark overall success if refinement worked
            yield f"data: {json.dumps({'type': 'final_report', 'report': final_report, 'sources': unique_final_sources})}\n\n"

        except Exception as e:
            # Log and yield general error for the whole process
            logging.error(
                f"Streaming analysis failed for '{analysis_title}': {e}", exc_info=True
            )
            stream_span.set_status(Status(StatusCode.ERROR, str(e)))
            stream_span.record_exception(e)
            yield f"data: {json.dumps({'type': 'error', 'message': f'Analysis failed: {str(e)}'})}\n\n"
            final_report_for_log = f"Error during generation: {str(e)}"

        finally:
            # Final logging for the entire process
            duration = time.time() - start_time
            stream_span.set_attribute("duration.ms", round(duration * 1000))
            await log_interaction(
                current_user.username,
                f"Report Generation: {analysis_title}",
                final_report_for_log[:10000],  # Log truncated report/error
                duration,
            )
            logging.info(
                f"Finished analysis stream '{analysis_title}' in {duration:.2f}s"
            )


# --- Load Analysis Prompts from JSON file ---
# This replaces the hard-coded dictionary, allowing you to edit prompts in prompts.json
prompts_dir = os.path.dirname(os.path.abspath(__file__))
analysis_prompts_config = {}
PROMPTS_FILE_PATH = os.path.join(prompts_dir, "prompts.json")

try:
    if os.path.exists(PROMPTS_FILE_PATH):
        with open(PROMPTS_FILE_PATH, "r", encoding="utf-8") as f:
            analysis_prompts_config = json.load(f)
        logging.info(
            f"Successfully loaded {len(analysis_prompts_config)} prompts from {PROMPTS_FILE_PATH}"
        )
    else:
        logging.warning(
            f"{PROMPTS_FILE_PATH} not found. No analysis prompts will be available."
        )
        analysis_prompts_config = {}  # Ensure it's an empty dict
except json.JSONDecodeError as e:
    logging.error(f"Error decoding {PROMPTS_FILE_PATH}: {e}. No prompts loaded.")
    analysis_prompts_config = {}  # Ensure it's an empty dict
except Exception as e:
    logging.error(f"Error loading {PROMPTS_FILE_PATH}: {e}", exc_info=True)
    analysis_prompts_config = {}  # Ensure it's an empty dict
# --- End of New Prompt Loading Logic ---


# --- ANALYSIS ENDPOINT (Enhanced version - uses new stream function) ---
# --- ANALYSIS ENDPOINT (Enhanced version - uses new stream function) ---


@app.post(
    "/analyze-document", summary="Analyze uploaded document with routing & streaming"
)
async def analyze_document_enhanced(
    analysis_title: str = Form(...),
    file: Optional[UploadFile] = File(None),  # File is optional
    model: str = Form("o3-mini"),  # Default to larger context model
    word_count: int = Form(2000),
    custom_prompt: Optional[str] = Form(None),
    current_user: User = Depends(auth_dependency),
):
    """
    Enhanced endpoint using smart routing and hybrid context stuffing.
    Streams results step-by-step and then provides a final synthesized report.
    """
    start_time = time.time()  # For logging only if needed here

    # --- Validate Inputs ---
    if not clients.synthesis_openai_client or not clients.planning_openai_client:
        raise HTTPException(status_code=503, detail="AI services not configured.")

    # Handle case where no file is provided
    if file:
        if not file.filename or not file.filename.lower().endswith(".docx"):
            raise HTTPException(
                status_code=400, detail="Invalid or missing .docx file."
            )

    model_deployment_id = (
        config.mda_deployment_id if model == "o3-mini" else config.deployment_id
    )
    logging.info(
        f"Using model deployment: {model_deployment_id} for analysis '{analysis_title}'"
    )
    # --- Extract Text and Perform Safety Check ---
    extracted_text = ""

    if file:
        try:
            logging.info(f"Extracting text from '{file.filename}'...")
            extracted_text = await extract_text_from_docx(file)
            if not extracted_text or not extracted_text.strip():
                raise HTTPException(
                    status_code=400, detail="Document appears empty or unreadable."
                )
            logging.info(f"Extracted {len(extracted_text)} characters from document.")

            # --- Content Safety Check on extracted text ---
            logging.info("Performing content safety check on document text...")
            document_safety = await analyze_content_safety(extracted_text)
            if not document_safety.is_safe:
                logging.warning(
                    f"Document content flagged: {document_safety.categories}"
                )
                raise HTTPException(
                    400,
                    f"Uploaded document content blocked for safety reasons: {', '.join(document_safety.categories)}",
                )
            logging.info("Document content safety check passed.")
        except HTTPException as http_exc:
            raise http_exc  # Re-raise validation/safety errors
        except Exception as e:
            logging.error(
                f"Error during file processing for {file.filename}: {e}", exc_info=True
            )
            raise HTTPException(
                status_code=500, detail=f"Error processing file: {str(e)}"
            )

    # Perform safety check on custom prompt if provided
    if custom_prompt:
        user_input_safety = await analyze_content_safety(custom_prompt.strip())
        if not user_input_safety.is_safe:
            raise HTTPException(
                400,
                f"Input blocked for safety: {', '.join(user_input_safety.categories)}",
            )

    # This logic now correctly uses the 'analysis_prompts_config' loaded at the top of the file
    prompt_template = (
        custom_prompt.strip()
        if custom_prompt and custom_prompt.strip()
        else analysis_prompts_config.get(analysis_title)
    )

    if not prompt_template:
        logging.error(
            f"Analysis title '{analysis_title}' not found in config and no custom prompt provided."
        )
        raise HTTPException(
            status_code=404,
            detail=f"Analysis prompt '{analysis_title}' not found or empty.",
        )

    # --- Start Streaming Response ---
    logging.info(
        f"Starting enhanced analysis stream for '{analysis_title}' by user {current_user.username}"
    )
    # Set query context for tracing downstream LLM calls using OTel context
    ctx = Context()
    ctx = set_value(USER_QUERY_KEY, f"Document Analysis: {analysis_title}", context=ctx)
    token_otel = attach(ctx)  # Attach context

    try:
        return StreamingResponse(
            stream_analysis_process_enhanced(  # Call the new generator function
                extracted_text=extracted_text,
                prompt_template=prompt_template,
                model_deployment_id=model_deployment_id,
                word_count=word_count,
                current_user=current_user,
                analysis_title=analysis_title,
            ),
            media_type="text/event-stream",
        )
    finally:
        detach(token_otel)  # Ensure context is detached


@app.post(
    "/mda-chat",
    response_model=ConversationalChatResponse,
    summary="Conversational chat for MDA Reviewer",
)
async def mda_chat(request: ChatRequest, current_user: User = Depends(auth_dependency)):
    # Uses the original handle_chat_request logic
    return await handle_chat_request(request, current_user)


@app.post(
    "/conversational-chat",
    response_model=ConversationalChatResponse,
    summary="General conversational chat",
)
async def conversational_chat(
    request: ChatRequest, current_user: User = Depends(auth_dependency)
):
    # Uses the original handle_chat_request logic
    return await handle_chat_request(request, current_user)


# ==============================================================================
#  UPDATED CHAT HANDLER (MAPS + DEEP RESEARCH + FINANCE + RAG)
# ==============================================================================

from pydantic import BaseModel
from typing import List, Optional, Any

# ==========================================
#  PYDANTIC MODELS (Must be defined before use)
# ==========================================

class ChatMessage(BaseModel):
    role: str
    content: str

class ConversationalChatRequest(BaseModel):
    question: str
    history: List[ChatMessage] = []

class ConversationalChatResponse(BaseModel):
    answer: str
    plan: List[str] = []
    sources: List[Any] = []
    source: Optional[str] = "azure-search" # identifying which tool was used


# ==============================================================================
#  UPDATED CHAT HANDLER (Fixed for Authentication Wrapper)
# ==============================================================================

# NOTE: Do NOT add @app.post here, because 'mda_chat' calls this function.
# ==============================================================================
#  UPDATED CHAT HANDLER (Final Fix for Pipeline Argument)
# ==============================================================================

# ==============================================================================
#  FINAL CHAT HANDLER (Hybrid RAG + Substantial Finance Data)
# ==============================================================================

# ==============================================================================
#  FINAL CHAT HANDLER (Hybrid RAG + Substantial Finance + Intelligent Links)
# ==============================================================================

async def handle_chat_request(request: ConversationalChatRequest, current_user=None):
    try:
        print("🔥 DEBUG: THE NEW CODE IS RUNNING! (Final Version) 🔥")

        # 1. SETUP & DEFAULTS
        user_id = "dev_user"
        if current_user:
            # Handle if current_user is a dict or an object
            user_id = current_user.get("id", "dev_user") if isinstance(current_user, dict) else str(current_user)
            
        q_lower = request.question.lower()

        # --- FEATURE 1: GOOGLE MAPS GENERATION ---
        # Triggers if query asks to "show map", "location of", etc.
        if "map" in q_lower and any(x in q_lower for x in ["show", "generate", "location", "view"]):
            loc_query = request.question
            for keyword in ["map of", "map for", "location of", "show me"]:
                if keyword in q_lower:
                    parts = q_lower.split(keyword)
                    if len(parts) > 1:
                        loc_query = parts[1].strip()
                        break
            
            map_image = generate_google_map(loc_query)
            return ConversationalChatResponse(
                answer=f"### 🗺️ Location Analysis: {loc_query.title()}\n{map_image}",
                source="google-maps"
            )

        # --- FEATURE 2: CHECK DEEP RESEARCH STATUS ---
        if any(x in q_lower for x in ["check status", "is it done", "job status", "research status"]):
            job = research_manager.get_job_status(user_id)
            if not job:
                return ConversationalChatResponse(answer="No research job found for your current session.")
            
            if job["status"] == "completed":
                return ConversationalChatResponse(
                    answer=f"### ✅ Research Complete\n\n{job['result']}",
                    source="deep-research-history"
                )
            elif job["status"] == "failed":
                return ConversationalChatResponse(answer=f"❌ Research Failed: {job.get('error')}")
            else:
                return ConversationalChatResponse(answer="🔄 Research In Progress. Please check back in a few minutes.")

        # --- FEATURE 3: START DEEP RESEARCH ---
        # Triggers for "deep dive", "web search", "deep agent"
        start_keywords = ["deep research", "deep dive", "deep agent", "web search", "internet"]
        if any(k in q_lower for k in start_keywords) and "status" not in q_lower:
            asyncio.create_task(research_manager.run_deep_research_task(user_id, request.question, ""))
            return ConversationalChatResponse(
                answer="### 🚀 MarketResearch Started\nI have launched the KPMG Market Research agent to analyze this request. This may take 2-5 minutes depending on complexity.\n\n**Next Step:** You can ask me *'Check status'* in a few minutes to see the report.",
                source="system"
            )

        # --- FEATURE 4: LIVE FINANCIAL INTELLIGENCE (PRO) ---
        # Fetches deep metrics (PE, Growth) using Python. No code interpreter.
        finance_keywords = [" live financial", " live market cap","live stock price" ]
        need_finance = any(k in q_lower for k in finance_keywords)
        
        finance_context_str = ""
        finance_visuals = ""
        
        if need_finance:
            print(f"🔍 DEBUG: Finance Intent Detected for: {q_lower}")
            
            # 1. Extract Tickers (e.g., "Tata" -> "TATAMOTORS.NS")
            tickers = await extract_tickers(request.question)
            
            if tickers:
                print(f"🔍 DEBUG: Extracted Tickers: {tickers}")
                
                # 2. Get "Substantial" Context (PE, Growth, Analyst Ratings)
                # This injects the data directly into the LLM's prompt.
                finance_context_str = get_llm_financial_context(tickers)
                
                # 3. Get ONLY the Markdown Table (No PNG generation needed)
                table_md = get_financial_table_markdown(tickers)
                
                finance_visuals = table_md

        # --- FEATURE 5: MAIN RAG PIPELINE (Standard Chat) ---
        # Inject the financial data (if any) into the user's query
        augmented_query = request.question
        if finance_context_str:
            augmented_query = f"{finance_context_str}\n\nUSER QUESTION: {request.question}"

        # Execute existing RAG pipeline (Azure Search + LLM Synthesis)
        rag_result = await execute_advanced_rag_pipeline(
            augmented_query, 
            history=request.history
        )
        
        # --- INTELLIGENT SOURCE FORMATTING (FIX) ---
        # Ensures Deep Research links remain clean URLs while PDFs get proxied
        final_sources = []
        for s in rag_result.get("sources", []):
            # 1. Check if it's a Clean Web URL (Deep Research)
            if isinstance(s, str) and (s.startswith("http://") or s.startswith("https://")):
                final_sources.append(s)
            
            # 2. Check for Finance Tool Tag
            elif s == "finance-tool":
                final_sources.append(s)
                
            # 3. Default: Assume it's an Internal File (Azure Blob)
            else:
                final_sources.append(generate_blob_proxy_url(s))

        # Combine the AI's analysis with the Data Table
        final_answer = rag_result["answer"] + finance_visuals

        return ConversationalChatResponse(
            answer=final_answer,
            plan=rag_result.get("plan", []),
            sources=final_sources, # Use the fixed source list
            source="hybrid-rag"
        )

    except Exception as e:
        logging.error(f"Chat Error: {e}", exc_info=True)
        return ConversationalChatResponse(answer=f"An error occurred while processing your request: {str(e)}")

@app.post(
    "/analyze-csv",
    response_model=CSVAnalysisResponse,
    summary="Analyze an uploaded CSV file with a code interpreter",
)
async def analyze_csv(
    prompt: str = Form(...),
    session_id: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
    current_user: User = Depends(auth_dependency),
):
    if not session_id:
        if not file:
            raise HTTPException(400, "A file must be uploaded for a new session.")
        session_id = str(uuid.uuid4())
        try:
            content = await file.read()
            # Run pandas read_csv in a thread
            df = await asyncio.to_thread(pd.read_csv, io.BytesIO(content))
            csv_sessions[session_id] = df
            logging.info(
                f"Started new CSV analysis session: {session_id} for user {current_user.username}"
            )
        except Exception as e:
            raise HTTPException(400, f"Failed to read CSV file: {e}")
    if session_id not in csv_sessions:
        raise HTTPException(404, "Session not found.")
    df = csv_sessions[session_id]

    code_generation_prompt = f"""You are a data analyst Python programmer. Given a pandas DataFrame named 'df' and a user prompt, generate Python code to perform the requested analysis. The code should print any textual output and can save a plot to 'plot.png' if visualization is needed.

User Prompt: {prompt}
DataFrame Info:
{df.info(verbose=False)}

Python Code:
"""
    try:

        @trace_openai_call(
            clients.synthesis_openai_client, config.deployment_id, "csv_code_generation"
        )
        async def make_csv_code_call():
            return await asyncio.to_thread(
                clients.synthesis_openai_client.chat.completions.create,
                model=config.deployment_id,
                messages=[{"role": "user", "content": code_generation_prompt}],
                temperature=0.0,
            )

        response = await make_csv_code_call()
        generated_code = (
            response.choices[0]
            .message.content.strip()
            .replace("```python", "")
            .replace("```", "")
        )

        def execute_code():
            import matplotlib.pyplot as plt
            from io import StringIO
            import sys

            old_stdout = sys.stdout
            sys.stdout = captured_output = StringIO()
            plot_path = "plot.png"
            if os.path.exists(plot_path):
                os.remove(plot_path)

            exec_globals = {"pd": pd, "df": df, "plt": plt}
            exec(generated_code, exec_globals)
            sys.stdout = old_stdout

            text_result = captured_output.getvalue()
            image_result = None
            if os.path.exists(plot_path):
                with open(plot_path, "rb") as f:
                    image_result = base64.b64encode(f.read()).decode("utf-8")
                os.remove(plot_path)
            return text_result, image_result

        text_result, image_result = await asyncio.to_thread(execute_code)
        return CSVAnalysisResponse(
            session_id=session_id, text_output=text_result, image_output=image_result
        )
    except Exception as e:
        logging.error(
            f"Error during code execution for session {session_id}: {e}", exc_info=True
        )
        return CSVAnalysisResponse(session_id=session_id, error=str(e))


@app.post(
    "/refine-report",
    response_model=RefineReportResponse,
    summary="Refine a report with new information",
)
async def refine_report(
    request: RefineReportRequest, current_user: User = Depends(auth_dependency)
):
    if not clients.synthesis_openai_client:
        raise HTTPException(503, "AI service not configured.")
    try:
        # 1. INPUT CHECK
        safety_result_input = await analyze_content_safety(request.new_info)
        if not safety_result_input.is_safe:
            logging.warning(
                f"Refine report new_info blocked for safety: {safety_result_input.categories}"
            )
            raise HTTPException(
                status_code=400,
                detail=f"Input blocked for safety: {', '.join(safety_result_input.categories)}",
            )

        prompt = f"""You are a report writing expert. Your task is to seamlessly integrate a new piece of information into an existing report. Do not simply append the new information. Instead, find the most relevant section in the 'ORIGINAL REPORT' and intelligently merge the 'NEW INFORMATION' into it. Rewrite paragraphs as needed to ensure the final report is coherent, clean, and well-integrated. Return ONLY the full, updated report text.
        
**ORIGINAL REPORT:**
---
{request.original_report}
---
**NEW INFORMATION TO INTEGRATE:**
---
{request.new_info}
---
**FULL, REFINED, AND INTEGRATED REPORT:**"""

        @trace_openai_call(
            clients.synthesis_openai_client, config.deployment_id, "refine_report"
        )
        async def make_refine_call():
            return await asyncio.to_thread(
                clients.synthesis_openai_client.chat.completions.create,
                model=config.deployment_id,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
                max_tokens=16000,
            )

        response = await make_refine_call()
        generated_report = response.choices[0].message.content

        # 2. OUTPUT CHECK
        safety_result_output = await analyze_content_safety(generated_report)
        if not safety_result_output.is_safe:
            logging.warning(
                f"AI output for refine_report blocked for safety: {safety_result_output.categories}"
            )
            safe_report = "The generated report was blocked for violating content safety policies."
            return RefineReportResponse(refined_report=safe_report)

        return RefineReportResponse(refined_report=generated_report)

    # --- THIS IS THE FIX ---
    except HTTPException as http_exc:
        # Re-raise HTTPExceptions (like our 400 safety block) directly
        raise http_exc
    except Exception as e:
        # Catch any other unexpected errors
        logging.error(f"Report refinement error: {e}", exc_info=True)
        raise HTTPException(500, f"Report refinement error: {e}")
    # --- END OF FIX ---


@app.post("/download-report", summary="Download a report as a .docx file")
async def download_report(
    request: DownloadRequest,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(auth_dependency),
):
    try:

        def create_doc():
            document = Document()
            styles = document.styles
            if "List Bullet" not in styles:
                styles.add_style("List Bullet", 1)
            if "List Number" not in styles:
                styles.add_style("List Number", 1)
            soup = BeautifulSoup(request.html_content, "html.parser")
            parse_html_to_docx(soup, document)

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                document.save(tmp.name)
                return tmp.name

        temp_file_path = await asyncio.to_thread(create_doc)
        background_tasks.add_task(os.remove, temp_file_path)
        return FileResponse(
            path=temp_file_path,
            filename="BIAL_Analysis_Report.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        logging.error("Error creating Word document", exc_info=True)
        raise HTTPException(500, f"Error creating Word document: {e}")


@app.get("/blob-proxy", summary="Proxy endpoint for Azure Blob Storage access")
async def blob_proxy(
    blob: str,
    request: Request,
    current_user: User = Depends(auth_dependency),
):
    """
    Secure proxy endpoint to access Azure Blob Storage files.
    This endpoint allows external users to access private blob storage
    without exposing direct blob URLs or SAS tokens in the frontend.

    Args:
        blob: Base64-encoded blob identifier in format "container/blob_name"
        request: FastAPI Request object for accessing headers

    Returns:
        StreamingResponse with the blob content
    """
    try:
        # Check if blob storage credentials are configured
        if (
            not config.azure_storage_account_name
            or not config.azure_storage_account_key
        ):
            raise HTTPException(
                status_code=503, detail="Azure Blob Storage credentials not configured"
            )

        # Decode the blob identifier
        try:
            # Add padding if needed for base64 decoding
            blob_identifier = blob + "=" * (4 - len(blob) % 4)
            decoded = base64.urlsafe_b64decode(blob_identifier.encode("utf-8")).decode(
                "utf-8"
            )
            container_name, blob_name = decoded.split("/", 1)
        except Exception as e:
            logging.error(f"Error decoding blob identifier: {e}", exc_info=True)
            raise HTTPException(
                status_code=400, detail="Invalid blob identifier format"
            )

        # Create BlobServiceClient
        connection_string = (
            f"DefaultEndpointsProtocol=https;"
            f"AccountName={config.azure_storage_account_name};"
            f"AccountKey={config.azure_storage_account_key};"
            f"EndpointSuffix=core.windows.net"
        )

        def get_blob_content():
            """Synchronous function to fetch blob content from Azure Storage."""
            try:
                blob_service_client = BlobServiceClient.from_connection_string(
                    connection_string
                )
                blob_client = blob_service_client.get_blob_client(
                    container=container_name, blob=blob_name
                )

                # Download blob content
                blob_data = blob_client.download_blob()
                content = blob_data.readall()

                # Get content type from blob properties
                properties = blob_client.get_blob_properties()
                content_type = (
                    properties.content_settings.content_type
                    or "application/octet-stream"
                )

                return content, content_type, blob_name
            except Exception as e:
                logging.error(
                    f"Error fetching blob {container_name}/{blob_name}: {e}",
                    exc_info=True,
                )
                raise

        # Fetch blob content asynchronously
        blob_content, content_type, filename = await asyncio.to_thread(get_blob_content)

        # Extract filename from blob name for Content-Disposition header
        filename_only = blob_name.split("/")[-1]

        # Create a streaming response
        def generate():
            yield blob_content

        return StreamingResponse(
            generate(),
            media_type=content_type,
            headers={
                "Content-Disposition": f'inline; filename="{filename_only}"',
                "Cache-Control": "private, max-age=3600",
                "X-Content-Type-Options": "nosniff",
            },
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Error in blob proxy endpoint: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error accessing blob: {str(e)}")


@app.post("/feedback", summary="Record user feedback on a response")
async def handle_feedback(
    request: FeedbackRequest, current_user: User = Depends(auth_dependency)
):
    if not clients.mongo_client:
        raise HTTPException(503, "DB not configured.")
    try:
        doc = request.dict()
        doc.update(
            {"timestamp": datetime.now(timezone.utc), "user": current_user.username}
        )
        await asyncio.to_thread(
            clients.mongo_client[config.cosmos_database_name][
                config.cosmos_feedback_collection
            ].insert_one,
            doc,
        )
        return {"status": "success", "message": "Feedback recorded."}
    except Exception as e:
        logging.error(f"Error recording feedback", exc_info=True)
        raise HTTPException(500, f"Failed to record feedback: {e}")


# --- NEW USER MANAGEMENT API ENDPOINTS ---
@app.post(
    "/auth/manage/create-user",
    response_model=NewUserResponse,
    summary="Create a new user",
    tags=["User Management"],
)
async def create_user(user: NewUserCreate, request: Request):
    require_api_secret(request)
    if not clients.mongo_client:
        raise HTTPException(503, "Database not configured.")
    if await get_user_from_db(user.username):
        raise HTTPException(400, "Username already exists.")

    hashed_password = get_password_hash(user.password)
    doc = {
        "id": f"user_{uuid.uuid4().hex[:8]}",
        "username": user.username,
        "email": user.email,
        "password_hash": hashed_password,
        "hashed_password": hashed_password,
        "first_name": user.first_name,
        "last_name": user.last_name,
        "role": user.role or "user",
        "permissions": [],
        "isActive": True,
        "createdAt": datetime.now(timezone.utc),
        "created_at": datetime.now(timezone.utc),
    }
    await asyncio.to_thread(
        clients.mongo_client[config.cosmos_database_name][
            config.cosmos_users_collection
        ].insert_one,
        doc,
    )
    return NewUserResponse(
        success=True,
        message="User created successfully",
        username=user.username,
        action="create",
        details={
            "user_id": doc["id"],
            "email": user.email,
            "first_name": user.first_name,
            "last_name": user.last_name,
            "role": user.role or "user",
            "isActive": True,
            "created_at": doc["created_at"].isoformat(),
        },
    )


@app.put(
    "/auth/manage/update-user",
    response_model=NewUserResponse,
    summary="Update user information",
    tags=["User Management"],
)
async def update_user(user_update: NewUserUpdate, request: Request):
    require_api_secret(request)
    if not clients.mongo_client:
        raise HTTPException(503, "Database not configured.")
    if not user_update.username:
        raise HTTPException(400, "Username is required.")
    if not await get_user_from_db(user_update.username):
        raise HTTPException(404, "User not found.")

    update_doc = {}
    if user_update.email is not None:
        update_doc["email"] = user_update.email
    if user_update.password is not None:
        hashed_password = get_password_hash(user_update.password)
        update_doc["password_hash"] = hashed_password
        update_doc["hashed_password"] = hashed_password
    if user_update.first_name is not None:
        update_doc["first_name"] = user_update.first_name
    if user_update.last_name is not None:
        update_doc["last_name"] = user_update.last_name
    if user_update.role is not None:
        update_doc["role"] = user_update.role
    if user_update.permissions is not None:
        update_doc["permissions"] = user_update.permissions

    if not update_doc:
        raise HTTPException(400, "No fields provided for update.")

    result = await asyncio.to_thread(
        clients.mongo_client[config.cosmos_database_name][
            config.cosmos_users_collection
        ].update_one,
        {"username": user_update.username},
        {"$set": update_doc},
    )
    if result.modified_count == 0:
        raise HTTPException(400, "No changes made to user.")

    updated_user = await get_user_from_db(user_update.username)
    if not updated_user:
        raise HTTPException(500, "Failed to retrieve updated user data.")

    return NewUserResponse(
        success=True,
        message="User updated successfully",
        username=updated_user.username,
        action="update",
        details={
            "user_id": getattr(updated_user, "id", "unknown"),
            "email": getattr(updated_user, "email", ""),
            "first_name": getattr(updated_user, "first_name", ""),
            "last_name": getattr(updated_user, "last_name", ""),
            "role": getattr(updated_user, "role", "user"),
            "permissions": getattr(updated_user, "permissions", {}),
            "isActive": getattr(updated_user, "isActive", True),
            "updated_at": datetime.now(timezone.utc).isoformat(),
        },
    )


@app.post(
    "/auth/manage/disable-user",
    response_model=NewUserResponse,
    summary="Disable or enable a user",
    tags=["User Management"],
)
async def disable_user(disable_request: NewUserDisableRequest, request: Request):
    require_api_secret(request)
    if not clients.mongo_client:
        raise HTTPException(503, "Database not configured.")
    if not await get_user_from_db(disable_request.username):
        raise HTTPException(404, "User not found.")

    isActive = not disable_request.disable
    result = await asyncio.to_thread(
        clients.mongo_client[config.cosmos_database_name][
            config.cosmos_users_collection
        ].update_one,
        {"username": disable_request.username},
        {"$set": {"isActive": isActive}},
    )

    if result.modified_count == 0:
        raise HTTPException(400, "No changes made to user status.")

    status_text = "disabled" if disable_request.disable else "enabled"
    return NewUserResponse(
        success=True,
        message=f"User {disable_request.username} has been {status_text}",
        username=disable_request.username,
        action=status_text,
        details={
            "isActive": isActive,
            "status": status_text,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        },
    )


@app.get(
    "/auth/manage/users",
    response_model=List[NewUserListResponse],
    summary="Get all users information",
    tags=["User Management"],
)
async def get_users(request: Request):
    require_api_secret(request)
    if not clients.mongo_client:
        raise HTTPException(503, "Database not configured.")

    def db_call():
        return list(
            clients.mongo_client[config.cosmos_database_name][
                config.cosmos_users_collection
            ].find({})
        )

    users_cursor = await asyncio.to_thread(db_call)
    users = []
    for user_doc in users_cursor:
        user_doc.setdefault("email", user_doc.get("username", ""))
        user_doc.setdefault("first_name", "")
        user_doc.setdefault("last_name", "")
        if "permissions" not in user_doc or isinstance(user_doc["permissions"], list):
            user_doc["permissions"] = {p: True for p in user_doc.get("permissions", [])}
        user_doc.setdefault(
            "created_at", user_doc.get("createdAt", datetime.now(timezone.utc))
        )
        if "id" not in user_doc and "_id" in user_doc:
            user_doc["id"] = str(user_doc["_id"])

        users.append(
            NewUserListResponse(
                id=user_doc["id"],
                username=user_doc["username"],
                email=user_doc.get("email", ""),
                first_name=user_doc.get("first_name", ""),
                last_name=user_doc.get("last_name", ""),
                role=user_doc.get("role", "user"),
                permissions=user_doc.get("permissions", {}),
                account_locked=not user_doc.get("isActive", True),
                created_at=user_doc.get("created_at"),
                last_login_at=user_doc.get("lastLogin"),
            )
        )
    return users


@app.get(
    "/auth/manage/roles",
    response_model=Dict[str, Any],
    summary="Get available roles and their details",
    tags=["User Management"],
)
async def get_role_details(request: Request):
    require_api_secret(request)
    roles = ["admin", "user", "manager", "viewer"]
    return {"success": True, "roles": roles}
